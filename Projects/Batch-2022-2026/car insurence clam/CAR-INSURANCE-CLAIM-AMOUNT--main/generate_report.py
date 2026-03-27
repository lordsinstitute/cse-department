#!/usr/bin/env python3
"""
Generate Major Project Report for Car Insurance Claim Amount Prediction Using Machine Learning
Based on A9 report format (matching exactly).
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ============================================================
# CONFIGURATION
# ============================================================
PROJECT_TITLE = "Car Insurance Claim Amount Prediction Using Machine Learning"
STUDENTS = [
    ("160922733142", "Shadman Ahmad"),
    ("160922733143", "Shaik Sufyaan Ahmed"),
    ("160922733147", "Syed Affan Hussain Syed Barey"),
    ("160922733151", "Syed Jawad"),
]
GUIDE_NAME = "Name of the Guide"
GUIDE_DESIGNATION = "Designation"
GUIDE_DEPT = "Dept. of CSE"
HOD_NAME = "Dr. TK Shaik Shavali"
PRINCIPAL_NAME = "Dr. Ravi Kishore Singh"
ACADEMIC_YEAR = "2024-2025"
YEAR = "2026"
COLLEGE = "LORDS INSTITUTE OF ENGINEERING AND TECHNOLOGY"
COLLEGE_SHORT = "LORDS INSTITUTE OF ENGINEERING & TECHNOLOGY"
DEPT = "Department of Computer Science & Engineering"
DEPT_FULL = "Department of Computer Science and Engineering"

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
LOGO_PATH = os.path.join(SCRIPT_DIR, "lords_logo.png")
FIGURES_DIR = os.path.join(SCRIPT_DIR, "figures")
SCREENSHOTS_DIR = os.path.join(SCRIPT_DIR, "screenshots")
OUTPUT_PATH = os.path.join(SCRIPT_DIR, "Car_Insurance_Claim_Prediction_ML_Major_Project_Report.docx")

# ============================================================
# DOCUMENT SETUP
# ============================================================
doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)


# ============================================================
# HELPER FUNCTIONS
# ============================================================
def add_centered_text(text, font_size=12, bold=False, color=None, space_after=6, space_before=0, keep_with_next=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if keep_with_next:
        p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = 'Times New Roman'
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p


def add_justified_text(text, font_size=12, bold=False, space_after=6, space_before=0, first_line_indent=None, keep_with_next=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.5
    if keep_with_next:
        p.paragraph_format.keep_with_next = True
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = 'Times New Roman'
    run.bold = bold
    return p


def add_left_text(text, font_size=12, bold=False, space_after=6, space_before=0, keep_with_next=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.5
    if keep_with_next:
        p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = 'Times New Roman'
    run.bold = bold
    return p


def add_chapter_heading(chapter_num, title):
    p1 = add_centered_text(f"CHAPTER {chapter_num}", font_size=18, bold=True, space_before=24, space_after=3)
    p1.paragraph_format.keep_with_next = True
    p2 = add_centered_text(title.upper(), font_size=16, bold=True, space_after=10)
    p2.paragraph_format.keep_with_next = True


def add_section_heading(number, title, font_size=16):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(f"{number}  {title}")
    run.font.size = Pt(font_size)
    run.font.name = 'Times New Roman'
    run.bold = True
    return p


def add_subsection_heading(number, title, font_size=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(f"{number}  {title}")
    run.font.size = Pt(font_size)
    run.font.name = 'Times New Roman'
    run.bold = True
    return p


def add_bullet(text, font_size=12):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.5
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = 'Times New Roman'
    return p


def set_cell_text(cell, text, bold=False, font_size=11, align=WD_ALIGN_PARAGRAPH.LEFT, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = 'Times New Roman'
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def shade_cell(cell, color="D9E2F3"):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def keep_table_on_one_page(table):
    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        cantSplit = parse_xml(f'<w:cantSplit {nsdecls("w")} w:val="true"/>')
        trPr.append(cantSplit)
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.keep_with_next = True
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcMar = parse_xml(
                    f'<w:tcMar {nsdecls("w")}>'
                    '<w:top w:w="30" w:type="dxa"/>'
                    '<w:bottom w:w="30" w:type="dxa"/>'
                    '</w:tcMar>'
                )
                existing = tcPr.findall(qn('w:tcMar'))
                for e in existing:
                    tcPr.remove(e)
                tcPr.append(tcMar)


def add_page_break():
    doc.add_page_break()


def add_figure(image_path, caption=None, width=Inches(5.0)):
    if os.path.exists(image_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture(image_path, width=width)
        p_img.paragraph_format.space_after = Pt(3)
    if caption:
        add_centered_text(caption, font_size=10, bold=True, space_after=8)


def add_letterhead_header(colored=False):
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    logo_cell = t.cell(0, 0)
    logo_cell.width = Inches(1.2)
    logo_para = logo_cell.paragraphs[0]
    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(LOGO_PATH):
        logo_para.add_run().add_picture(LOGO_PATH, width=Inches(1.0))
    text_cell = t.cell(0, 1)
    text_cell.width = Inches(5.0)
    p = text_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(COLLEGE_SHORT)
    run.font.size = Pt(13)
    run.font.name = 'Times New Roman'
    run.bold = True
    if colored:
        run.font.color.rgb = RGBColor(255, 0, 0)
    p2 = text_cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("(UGC Autonomous)")
    r2.font.size = Pt(10)
    r2.font.name = 'Times New Roman'
    p3 = text_cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Approved by AICTE | Affiliated to Osmania University | Estd.2003.")
    r3.font.size = Pt(9)
    r3.font.name = 'Times New Roman'
    p4 = text_cell.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run("Accredited with \u2018A\u2019 grade by NAAC | Accredited by NBA")
    r4.font.size = Pt(9)
    r4.font.name = 'Times New Roman'
    p5 = text_cell.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r5 = p5.add_run(DEPT)
    r5.font.size = Pt(12)
    r5.font.name = 'Times New Roman'
    r5.bold = True
    if colored:
        r5.font.color.rgb = RGBColor(0, 128, 0)
    for cell in [logo_cell, text_cell]:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '</w:tcBorders>')
        tcPr.append(tcBorders)
    return t


def add_page_number(section_obj, start=1, fmt='decimal'):
    sectPr = section_obj._sectPr
    pgNumType = parse_xml(f'<w:pgNumType {nsdecls("w")} w:start="{start}" w:fmt="{fmt}"/>')
    existing = sectPr.findall(qn('w:pgNumType'))
    for e in existing:
        sectPr.remove(e)
    sectPr.append(pgNumType)


# ============================================================
# PAGE i - TITLE PAGE
# ============================================================
add_centered_text("A", font_size=14, space_before=12, space_after=0)
add_centered_text("Major Project Report", font_size=16, bold=True, space_after=4)
add_centered_text("on", font_size=12, space_after=4)
add_centered_text(PROJECT_TITLE, font_size=18, bold=True, color=(255, 0, 0), space_after=6)
add_centered_text("submitted in partial fulfillment of the requirement for the award of the degree of",
                   font_size=11, space_after=4)
add_centered_text("BACHELOR OF ENGINEERING", font_size=13, bold=True, space_after=2)
add_centered_text("In", font_size=12, space_after=2)
add_centered_text("COMPUTER SCIENCE & ENGINEERING", font_size=13, bold=True, space_after=6)
add_centered_text("By", font_size=12, space_after=4)

t = doc.add_table(rows=len(STUDENTS), cols=2)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (roll, name) in enumerate(STUDENTS):
    set_cell_text(t.cell(i, 0), name, font_size=12, bold=True)
    set_cell_text(t.cell(i, 1), roll, font_size=12, bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    t.cell(i, 0).width = Inches(3)
    t.cell(i, 1).width = Inches(2.5)

add_centered_text("", space_after=1)
add_centered_text("Under the esteemed guidance of", font_size=12, space_after=2)
add_centered_text(GUIDE_NAME, font_size=12, bold=True, space_after=2)
add_centered_text(f"{GUIDE_DESIGNATION} & {GUIDE_DEPT}", font_size=12, space_after=6)

p_logo = doc.add_paragraph()
p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
if os.path.exists(LOGO_PATH):
    p_logo.add_run().add_picture(LOGO_PATH, width=Inches(1.3))

add_centered_text(DEPT, font_size=14, bold=True, space_before=4, space_after=3)
add_centered_text(COLLEGE, font_size=13, bold=True, color=(255, 0, 0), space_after=2)
add_centered_text("(UGC Autonomous)", font_size=11, space_after=1)
add_centered_text("Approved by AICTE | Affiliated to Osmania University | Estd.2003", font_size=10, space_after=1)
add_centered_text("Sy.No.32, Himayat Sagar, Near TGPA Junction, Hyderabad-500091, India.", font_size=10, space_after=3)
add_centered_text(f"({YEAR})", font_size=14, bold=True, space_after=3)

add_page_number(doc.sections[0], start=1, fmt='lowerRoman')

# ============================================================
# PAGE ii - CERTIFICATE
# ============================================================
add_letterhead_header()
add_centered_text("", space_after=2)
add_centered_text("CERTIFICATE", font_size=16, bold=True, space_after=8)

cert_p = doc.add_paragraph()
cert_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
cert_p.paragraph_format.space_after = Pt(6)
cert_p.paragraph_format.line_spacing = 1.5
cert_p.paragraph_format.first_line_indent = Cm(1.27)
r = cert_p.add_run(f'This is to certify that the project report entitled \u201c{PROJECT_TITLE}\u201d being Submitted by ')
r.font.size = Pt(12)
r.font.name = 'Times New Roman'
student_str = ", ".join(f"{name} ({roll})" for roll, name in STUDENTS[:-1]) + f", and {STUDENTS[-1][1]} ({STUDENTS[-1][0]})" if len(STUDENTS) > 1 else f"{STUDENTS[0][1]} ({STUDENTS[0][0]})"
r3 = cert_p.add_run(student_str)
r3.font.size = Pt(12)
r3.font.name = 'Times New Roman'
r3.bold = True
r5 = cert_p.add_run(
    f' in partial fulfillment of the requirements for the award of '
    f'the degree of Bachelor of Engineering in Computer Science and Engineering during the '
    f'academic year {ACADEMIC_YEAR}.'
)
r5.font.size = Pt(12)
r5.font.name = 'Times New Roman'
add_justified_text(
    "This is further certified that the work done under my guidance, and the results of this work "
    "have not been submitted elsewhere for the award of any of the degree",
    first_line_indent=1.27, space_after=18
)

sig = doc.add_table(rows=2, cols=2)
sig.alignment = WD_TABLE_ALIGNMENT.CENTER
set_cell_text(sig.cell(0, 0), "Internal Guide", bold=True, font_size=11)
set_cell_text(sig.cell(0, 1), "Head of the Department", bold=True, font_size=11, align=WD_ALIGN_PARAGRAPH.RIGHT)
set_cell_text(sig.cell(1, 0), f"{GUIDE_NAME}\n{GUIDE_DESIGNATION}", font_size=11)
set_cell_text(sig.cell(1, 1), f"{HOD_NAME}\nHOD - CSE", font_size=11, align=WD_ALIGN_PARAGRAPH.RIGHT)

add_centered_text("", space_after=12)

sig2 = doc.add_table(rows=2, cols=2)
sig2.alignment = WD_TABLE_ALIGNMENT.CENTER
set_cell_text(sig2.cell(0, 0), "Principal", bold=True, font_size=11)
set_cell_text(sig2.cell(0, 1), "External Examiner", bold=True, font_size=11, align=WD_ALIGN_PARAGRAPH.RIGHT)
set_cell_text(sig2.cell(1, 0), PRINCIPAL_NAME, font_size=11)
set_cell_text(sig2.cell(1, 1), "Date:", font_size=11, align=WD_ALIGN_PARAGRAPH.RIGHT)

# ============================================================
# PAGE iii - DECLARATION
# ============================================================
add_page_break()
add_letterhead_header(colored=True)
add_centered_text("", space_after=6)
add_centered_text("DECLARATION BY THE CANDIDATE", font_size=16, bold=True, space_after=12,
                   color=(0x1F, 0x4E, 0x79))

decl_text = (
    f'We, hereby declare that the project report entitled \u201c{PROJECT_TITLE}\u201d, under '
    f'the guidance of {GUIDE_NAME}, {GUIDE_DESIGNATION}, {DEPT_FULL}, '
    f'Lords Institute of Engineering & Technology, affiliated to Osmania University, Hyderabad '
    f'is submitted in partial fulfillment of the requirements for the award of the degree of '
    f'Bachelor of Engineering in Computer Science and Engineering.'
)
add_justified_text(decl_text, first_line_indent=1.27)
add_justified_text(
    "This is a record of bonafide work carried out by us and the results of this work "
    "have not been reproduced or copied from any source. The results embodied in this project report "
    "have not been submitted to any other university or institute for the award of any other degree.",
    first_line_indent=1.27, space_after=24
)

t3 = doc.add_table(rows=len(STUDENTS), cols=2)
t3.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (roll, name) in enumerate(STUDENTS):
    set_cell_text(t3.cell(i, 0), name, font_size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(t3.cell(i, 1), roll, font_size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

# ============================================================
# PAGE iv - ACKNOWLEDGMENT
# ============================================================
add_page_break()
add_letterhead_header(colored=True)
add_centered_text("", space_after=6)
add_centered_text("ACKNOWLEDGMENT", font_size=16, bold=True, space_after=12,
                   color=(0x1F, 0x4E, 0x79))

add_justified_text(
    "First, we wish to thank GOD Almighty who created heavens and earth, who helped us in "
    "completing this project and we also thank our Parents who encouraged us in this period.",
    space_after=6
)
add_justified_text(
    f"We would like to thank {GUIDE_NAME}, {GUIDE_DESIGNATION}, {GUIDE_DEPT}, "
    f"Lords Institute of Engineering & Technology, affiliated to Osmania University, Hyderabad, "
    f"our project internal guide, for her guidance and help. Her insight during the course of our "
    f"major project and regular guidance were invaluable to us.",
    space_after=6
)
add_justified_text(
    f"We would like to express our deep sense of gratitude to {HOD_NAME}, Professor & "
    f"Head of the Department, Computer Science & Engineering, Lords Institute of Engineering "
    f"& Technology, affiliated to Osmania University, Hyderabad, for his encouragement and "
    f"cooperation throughout the project.",
    space_after=6
)
add_justified_text(
    f"We would also like to thank {PRINCIPAL_NAME}, Principal of our college, for extending his help.",
    space_after=18
)

t4 = doc.add_table(rows=len(STUDENTS)+1, cols=2)
t4.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (roll, name) in enumerate(STUDENTS):
    set_cell_text(t4.cell(i, 0), name, font_size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(t4.cell(i, 1), roll, font_size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_text(t4.cell(len(STUDENTS), 0), "(Lords Institute of Engineering and Technology)",
              font_size=11, align=WD_ALIGN_PARAGRAPH.CENTER)

# ============================================================
# PAGE v - VISION & MISSION OF THE INSTITUTE
# ============================================================
add_page_break()
add_letterhead_header(colored=True)
add_centered_text("", space_after=4)

add_left_text("Vision of the Institute:", bold=True, space_after=4)
add_justified_text(
    "Lords Institute of Engineering and Technology strives for excellence in professional education through "
    "quality, innovation and teamwork and aims to emerge as a premier institute in the state and across the nation.",
    first_line_indent=1.27, space_after=6
)

add_left_text("Mission of the Institute:", bold=True, space_after=4)
for m in [
    "To impart quality professional education that meets the needs of present and emerging technological world.",
    "To strive for student achievement and success, preparing them for life, career and leadership.",
    "To provide a scholarly and vibrant learning environment that enables faculty, staff and students to achieve personal and professional growth.",
    "To contribute to advancement of knowledge, in both fundamental and applied areas of engineering and technology.",
    "To forge mutually beneficial relationships with government organizations, industries, society and the alumni.",
]:
    add_bullet(m)

# ============================================================
# PAGE vi - VISION & MISSION OF THE DEPARTMENT
# ============================================================
add_page_break()
add_letterhead_header(colored=True)
add_centered_text("", space_after=4)

add_left_text("Vision of the Department:", bold=True, space_after=4)
add_justified_text(
    "To emerge as a center of excellence for quality Computer Science and Engineering education "
    "with innovation, leadership and values.",
    first_line_indent=1.27, space_after=6
)

add_left_text("Mission of the Department:", bold=True, space_after=4)
for dm in [
    "DM1: Provide fundamental and practical training through learner \u2013 centric Teaching-Learning Process and state-of-the-art infrastructure.",
    "DM2: Develop design, research, and entrepreneurial skills for successful career.",
    "DM3: Promote training and activities through Industry-Academia interactions.",
]:
    add_bullet(dm)
add_left_text("Note: DM: Department Mission", font_size=11, space_before=6, space_after=6)

# ============================================================
# PAGE vii - PEOs
# ============================================================
add_page_break()
add_letterhead_header(colored=True)
add_centered_text("", space_after=2)
add_centered_text("B.E. Computer Science and Engineering Program Educational Objectives (PEOs):",
                   font_size=12, bold=True, space_after=6, keep_with_next=True)

peos = [
    ("PEO1", "Exhibit strong foundations in Basic Sciences, Computer Science and allied engineering"),
    ("PEO2", "Identify, formulate, analyze and create professional solutions, novel products using appropriate tools and techniques, design skills."),
    ("PEO3", "Pursue successful career with continuous learning, with emphasis on competency oriented towards industry"),
    ("PEO4", "Practice ethics, managerial and leadership skills to work cohesively within a group"),
]
peo_table = doc.add_table(rows=4, cols=2)
peo_table.style = 'Table Grid'
peo_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (code, desc) in enumerate(peos):
    set_cell_text(peo_table.cell(i, 0), code, bold=True, font_size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(peo_table.cell(i, 1), desc, font_size=11)
    shade_cell(peo_table.cell(i, 0), "D9E2F3")
for row in peo_table.rows:
    row.cells[0].width = Inches(0.8)
    row.cells[1].width = Inches(5.4)
keep_table_on_one_page(peo_table)

# ============================================================
# PAGE viii - POs
# ============================================================
add_page_break()
add_letterhead_header(colored=True)
add_centered_text("", space_after=4)
add_centered_text("B.E. Computer Science and Engineering Program Outcomes (POs):", font_size=12, bold=True, space_after=3, keep_with_next=True)
add_left_text("Engineering Graduates will be able to:", font_size=12, space_after=4, keep_with_next=True)

po_table = doc.add_table(rows=12, cols=2)
po_table.style = 'Table Grid'
po_table.alignment = WD_TABLE_ALIGNMENT.CENTER
pos_data = [
    ("S. No.", "Program Outcomes (POs)"),
    ("1.", "PO1: Engineering Knowledge: Apply knowledge of mathematics, natural science, computing, engineering fundamentals and an engineering specialization as specified in WK1 to WK4 respectively to develop to the solution of complex engineering problems."),
    ("2.", "PO2: Problem Analysis: Identify, formulate, review research literature and analyze complex engineering problems reaching substantiated conclusions with consideration for sustainable development. (WK1 to WK4)"),
    ("3.", "PO3: Design/Development of Solutions: Design creative solutions for complex engineering problems and design/develop systems/components/processes to meet identified needs with consideration for the public health and safety, whole-life cost, net zero carbon, culture, society and environment as required. (WK5)"),
    ("4.", "PO4: Conduct Investigations of Complex Problems: Conduct investigations of complex engineering problems using research-based knowledge including design of experiments, modelling, analysis & interpretation of data to provide valid conclusions. (WK8)."),
    ("5.", "PO5: Engineering Tool Usage: Create, select and apply appropriate techniques, resources and modern engineering & IT tools, including prediction and modelling recognizing their limitations to solve complex engineering problems. (WK2 and WK6)"),
    ("6.", "PO6: The Engineer and The World: Analyze and evaluate societal and environmental aspects while solving complex engineering problems for its impact on sustainability with reference to economy, health, safety, legal framework, culture and environment. (WK1, WK5, and WK7)."),
    ("7.", "PO7: Ethics: Apply ethical principles and commit to professional ethics, human values, diversity and inclusion; adhere to national & international laws. (WK9)"),
    ("8.", "PO8: Individual and Collaborative Team work: Function effectively as an individual, and as a member or leader in diverse/multi-disciplinary teams."),
    ("9.", "PO9: Communication: Communicate effectively and inclusively within the engineering community and society at large, such as being able to comprehend and write effective reports and design documentation, make effective presentations considering cultural, language, and learning differences"),
    ("10.", "PO10: Project Management and Finance: Apply knowledge and understanding of engineering management principles and economic decision-making and apply these to one\u2019s own work, as a member and leader in a team, and to manage projects and in multidisciplinary environments."),
    ("11.", "PO11: Life-Long Learning: Recognize the need for, and have the preparation and ability for i) independent and life-long learning ii) adaptability to new and emerging technologies and iii) critical thinking in the broadest context of technological change. (WK8)"),
]
for i, (num, text) in enumerate(pos_data):
    set_cell_text(po_table.cell(i, 0), num, bold=(i == 0), font_size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(po_table.cell(i, 1), text, bold=(i == 0), font_size=10)
    if i == 0:
        shade_cell(po_table.cell(i, 0), "D9E2F3")
        shade_cell(po_table.cell(i, 1), "D9E2F3")
for row in po_table.rows:
    row.cells[0].width = Inches(0.6)
    row.cells[1].width = Inches(5.6)
keep_table_on_one_page(po_table)

# ============================================================
# PAGE ix - PSOs
# ============================================================
add_page_break()
add_letterhead_header(colored=True)
add_centered_text("", space_after=4)
add_centered_text("B.E. Computer Science and Engineering Program Specific Outcomes (PSO\u2019s):",
                   font_size=12, bold=True, space_after=6, keep_with_next=True)

pso_table = doc.add_table(rows=2, cols=2)
pso_table.style = 'Table Grid'
pso_table.alignment = WD_TABLE_ALIGNMENT.CENTER
psos = [
    ("PSO1", "Professional Skills:\u00a0Implement computer programs in the areas related to algorithms, system software, multimedia, web design, big data analytics and networking for efficient analysis and design of computer-based systems of varying complexity"),
    ("PSO2", "Problem-Solving Skills:\u00a0Apply standard practices and strategies in software service management using open-ended programming environment with agility to deliver a quality service for business success"),
]
for i, (code, desc) in enumerate(psos):
    set_cell_text(pso_table.cell(i, 0), code, bold=True, font_size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(pso_table.cell(i, 1), desc, font_size=11)
    shade_cell(pso_table.cell(i, 0), "D9E2F3")
for row in pso_table.rows:
    row.cells[0].width = Inches(0.8)
    row.cells[1].width = Inches(5.4)
keep_table_on_one_page(pso_table)

# ============================================================
# PAGE x - COURSE OUTCOMES
# ============================================================
add_page_break()
add_letterhead_header(colored=True)
add_centered_text("", space_after=4)
add_centered_text("Course Outcomes: C424 - Major Project", font_size=12, bold=True, space_after=2, keep_with_next=True)
add_left_text("Student will be able to", font_size=12, space_after=4, keep_with_next=True)

co_table = doc.add_table(rows=6, cols=3)
co_table.style = 'Table Grid'
co_table.alignment = WD_TABLE_ALIGNMENT.CENTER
cos = [
    ("CO. No", "Description", "Blooms\nTaxonomy\nLevel"),
    ("C424.1", "Demonstrate the ability to synthesize and apply the knowledge and skills acquired in the academic program to real-world problems.", "BTL3"),
    ("C424.2", "Evaluate different solutions based on economic and technical feasibility.", "BTL5"),
    ("C424.3", "Effectively plan a a project and confidently perform all aspects of project management", "BTL4"),
    ("C424.4", "Demonstrate effective oral communication skills", "BTL2"),
    ("C424.5", "Demonstrate effective team work and written skills", "BTL1"),
]
for i, (co, desc, bloom) in enumerate(cos):
    set_cell_text(co_table.cell(i, 0), co, bold=(i == 0), font_size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(co_table.cell(i, 1), desc, bold=(i == 0), font_size=10)
    set_cell_text(co_table.cell(i, 2), bloom, bold=(i == 0), font_size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    if i == 0:
        shade_cell(co_table.cell(i, 0), "D9E2F3")
        shade_cell(co_table.cell(i, 1), "D9E2F3")
        shade_cell(co_table.cell(i, 2), "D9E2F3")
for row in co_table.rows:
    row.cells[0].width = Inches(0.8)
    row.cells[1].width = Inches(4.2)
    row.cells[2].width = Inches(1.2)
keep_table_on_one_page(co_table)

# ============================================================
# PAGE xi - COURSE ARTICULATION MATRIX
# ============================================================
add_page_break()
add_letterhead_header(colored=True)
add_centered_text("", space_after=2)
add_centered_text("Course Articulation Matrix:", font_size=12, bold=True, space_after=1, keep_with_next=True)
add_centered_text("Mapping of Course Outcomes (CO) with Program Outcomes (PO) and Program Specific Outcomes (PSO\u2019s):",
                   font_size=11, space_after=2, keep_with_next=True)

cols = ["Course\nOutcome s\n(CO)", "PO1", "PO2", "PO3", "PO4", "PO5", "PO6", "PO7", "PO 8", "PO9", "PO10", "PO11", "PSO1", "PSO2"]
cam_table = doc.add_table(rows=8, cols=14)
cam_table.style = 'Table Grid'
cam_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, col_name in enumerate(cols):
    set_cell_text(cam_table.cell(0, j), col_name, bold=True, font_size=7, align=WD_ALIGN_PARAGRAPH.CENTER)
    shade_cell(cam_table.cell(0, j), "D9E2F3")
co_matrix = [
    ("C424.1.", [3, 3, 3, 1, 3, 2, 1, 1, 2, 2, 3, 3, 3]),
    ("C424.2.", [3, 3, 3, 2, 2, 3, 1, 1, 2, 3, 3, 3, 3]),
    ("C424.3.", [3, 3, 3, 3, 3, 2, 1, 1, 3, 2, 3, 3, 3]),
    ("C424.4.", [3, 2, 1, 1, 2, 2, 1, 2, 3, 2, 3, 3, 3]),
    ("C424.5.", [3, 1, 2, 1, 2, 2, 1, 2, 3, 2, 3, 3, 3]),
    ("Average", [3.0, 2.4, 2.4, 1.6, 2.4, 2.2, 1.0, 1.4, 2.6, 2.2, 3.0, 3.0, 3.0]),
]
for i, (label, vals) in enumerate(co_matrix):
    set_cell_text(cam_table.cell(i + 1, 0), label, font_size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    for j, v in enumerate(vals):
        text = str(v) if isinstance(v, int) else f"{v:.1f}"
        set_cell_text(cam_table.cell(i + 1, j + 1), text, font_size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
keep_table_on_one_page(cam_table)

add_left_text("Level:", font_size=10, bold=True, space_after=0, space_before=2)
add_left_text("1- Low correlation (Low), 2- Medium correlation (Medium), 3-High correlation (High)", font_size=10, space_after=3)

# ============================================================
# SDG MAPPING
# ============================================================
add_left_text("SDG Mapping:", font_size=12, bold=True, space_after=2, keep_with_next=True)
sdg_table = doc.add_table(rows=7, cols=6)
sdg_table.style = 'Table Grid'
sdg_table.alignment = WD_TABLE_ALIGNMENT.CENTER
sdg_headers = ["SDG", "Mapped\nIndicator", "SDG", "Mapped\nIndicator", "SDG", "Mapped\nIndicator"]
for j, h in enumerate(sdg_headers):
    set_cell_text(sdg_table.cell(0, j), h, bold=True, font_size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    shade_cell(sdg_table.cell(0, j), "D9E2F3")

SDG_COLORS = {
    1: "E5243B", 2: "DDA63A", 3: "4C9F38", 4: "C5192D",
    5: "FF3A21", 6: "26BDE2", 7: "FCC30B", 8: "A21942",
    9: "FD6925", 10: "DD1367", 11: "FD9D24", 12: "BF8B2E",
    13: "3F7E44", 14: "0A97D9", 15: "56C02B", 16: "00689D",
    17: "19486A",
}
sdg_data = [
    [(1, "1 NO\nPOVERTY", False),       (7, "7 AFFORDABLE AND\nCLEAN ENERGY", False),   (13, "13 CLIMATE\nACTION", False)],
    [(2, "2 ZERO\nHUNGER", False),      (8, "8 DECENT WORK AND\nECONOMIC GROWTH", True),(14, "14 LIFE\nBELOW WATER", False)],
    [(3, "3 GOOD HEALTH\nAND WELL-BEING", False), (9, "9 INDUSTRY, INNOVATION\nAND INFRASTRUCTURE", True), (15, "15 LIFE\nON LAND", False)],
    [(4, "4 QUALITY\nEDUCATION", False),  (10, "10 REDUCED\nINEQUALITIES", False),        (16, "16 PEACE, JUSTICE\nAND STRONG INSTITUTIONS", True)],
    [(5, "5 GENDER\nEQUALITY", False),   (11, "11 SUSTAINABLE CITIES\nAND COMMUNITIES", False), (17, "17 PARTNERSHIPS\nFOR THE GOALS", False)],
    [(6, "6 CLEAN WATER\nAND SANITATION", False), (12, "12 RESPONSIBLE\nCONSUMPTION AND PRODUCTION", False), (0, "", False)],
]
for i, row in enumerate(sdg_data):
    for k, (sdg_num, sdg_name, mapped) in enumerate(row):
        sdg_col = k * 2
        ind_col = k * 2 + 1
        if sdg_num > 0:
            set_cell_text(sdg_table.cell(i + 1, sdg_col), sdg_name, bold=True, font_size=8,
                          align=WD_ALIGN_PARAGRAPH.CENTER, color=(255, 255, 255))
            shade_cell(sdg_table.cell(i + 1, sdg_col), SDG_COLORS[sdg_num])
            set_cell_text(sdg_table.cell(i + 1, ind_col), "\u2713" if mapped else "", font_size=12,
                          align=WD_ALIGN_PARAGRAPH.CENTER)
        else:
            set_cell_text(sdg_table.cell(i + 1, sdg_col), "", font_size=9)
            set_cell_text(sdg_table.cell(i + 1, ind_col), "", font_size=9)
for row in sdg_table.rows:
    row.cells[0].width = Inches(1.3)
    row.cells[1].width = Inches(0.7)
    row.cells[2].width = Inches(1.3)
    row.cells[3].width = Inches(0.7)
    row.cells[4].width = Inches(1.3)
    row.cells[5].width = Inches(0.7)
keep_table_on_one_page(sdg_table)

# ============================================================
# ABSTRACT
# ============================================================
add_page_break()
add_centered_text("ABSTRACT", font_size=16, bold=True, space_before=24, space_after=12)

add_justified_text(
    "Car insurance claim prediction is a critical task in the insurance industry, enabling companies to "
    "assess risk, optimize premium pricing, and detect potentially fraudulent claims. Understanding the "
    "factors that influence whether a policyholder will file a claim is essential for actuaries, underwriters, "
    "and insurance managers to design effective risk management strategies and allocate resources efficiently. "
    "The insurance sector generates vast amounts of policyholder data, providing a rich dataset for statistical "
    "and machine learning analysis.",
    first_line_indent=1.27
)
add_justified_text(
    "This project presents a web-based machine learning platform for predicting car insurance claim outcomes "
    "using four classification algorithms: Random Forest, Gradient Boosting, Support Vector Machine (SVM), "
    "and Logistic Regression. The system processes a synthetic Car Insurance Claim dataset containing 10,000 "
    "records across 17 features including age, gender, driving experience, credit score, vehicle type, annual "
    "mileage, speeding violations, DUIs, and past accidents. Data preprocessing includes label encoding of "
    "9 categorical features, standard scaling of 8 numeric features, and an 80/20 stratified train-test split.",
    first_line_indent=1.27
)
add_justified_text(
    "The Gradient Boosting model achieved the highest performance with an accuracy of 91.95%, precision of "
    "89.45%, recall of 78.27%, and F1-score of 83.49%, followed by SVM (accuracy=91.10%), Random Forest "
    "(accuracy=90.10%), and Logistic Regression (accuracy=90.25%). The application provides Chart.js "
    "visualizations including distribution charts, feature importance plots, model comparison bar charts, "
    "and confusion matrices for comprehensive model evaluation.",
    first_line_indent=1.27
)
add_justified_text(
    "The platform is built using Flask (Python) with a Bootstrap 5 dark-themed user interface featuring "
    "responsive design and interactive dashboards. The system includes user authentication with SQLite "
    "database (Werkzeug password hashing), prediction history tracking, and exploratory data analysis "
    "visualization galleries. Users register, log in, input policyholder details, and receive claim "
    "predictions with confidence scores. The application is containerized with Docker for reproducible "
    "deployment on port 5002.",
    first_line_indent=1.27
)
add_justified_text(
    "Keywords: Car Insurance Claim Prediction, Machine Learning, Classification, Gradient Boosting, "
    "Random Forest, SVM, Logistic Regression, Flask, Bootstrap 5, SQLite, Chart.js, Docker.",
    first_line_indent=1.27, bold=True
)

# ============================================================
# TABLE OF CONTENTS
# ============================================================
add_page_break()
add_centered_text("TABLE OF CONTENTS", font_size=16, bold=True, space_before=24, space_after=12)

toc_entries = [
    ("Title Page", "i"),
    ("Certificate", "ii"),
    ("Declaration", "iii"),
    ("Acknowledgment", "iv"),
    ("Vision & Mission of the Institute", "v"),
    ("Vision & Mission of the Department", "vi"),
    ("Program Educational Objectives (PEOs)", "vii"),
    ("Program Outcomes (POs)", "viii"),
    ("Program Specific Outcomes (PSOs)", "ix"),
    ("Course Outcomes", "x"),
    ("Course Articulation Matrix", "xi"),
    ("SDG Mapping", "xi"),
    ("Abstract", "xii"),
    ("Table of Contents", "xiii"),
    ("List of Figures", "xiv"),
    ("List of Tables", "xv"),
    ("", ""),
    ("CHAPTER 1: INTRODUCTION", "1"),
    ("1.1    Introduction", "1"),
    ("1.2    Scope of the Project", "2"),
    ("1.3    Objectives", "3"),
    ("1.4    Problem Formulation", "3"),
    ("1.5    Existing System", "4"),
    ("1.6    Proposed System", "5"),
    ("", ""),
    ("CHAPTER 2: LITERATURE SURVEY", "7"),
    ("2.1 \u2013 2.15  Literature Reviews", "7"),
    ("", ""),
    ("CHAPTER 3: REQUIREMENT ANALYSIS AND SYSTEM SPECIFICATION", "15"),
    ("3.1    Feasibility Study", "15"),
    ("3.2    Software Requirement Specification", "16"),
    ("3.2.1    Overall Description", "16"),
    ("3.2.2    System Feature Requirement", "17"),
    ("3.2.3    Non-Functional Requirement", "18"),
    ("3.3    System Requirements", "19"),
    ("3.4    SDLC Model to be Used", "19"),
    ("3.5    Software Requirements", "20"),
    ("", ""),
    ("CHAPTER 4: SYSTEM DESIGN", "21"),
    ("4.1    Design Approach", "21"),
    ("4.2    System Architecture Diagram", "22"),
    ("4.3    UML Diagrams", "23"),
    ("4.4    User Interface Design", "26"),
    ("4.5    Database Schema", "27"),
    ("", ""),
    ("CHAPTER 5: IMPLEMENTATION", "29"),
    ("5.1    Methodologies", "29"),
    ("5.2    Implementation Details", "31"),
    ("5.3    Module Description", "32"),
    ("5.4    Sample Code", "33"),
    ("", ""),
    ("CHAPTER 6: TESTING", "37"),
    ("6.1    Types of Testing", "37"),
    ("6.2    Test Cases", "39"),
    ("", ""),
    ("CHAPTER 7: RESULTS AND DISCUSSION", "42"),
    ("7.1 \u2013 7.13  Application Screenshots", "42"),
    ("7.14 \u2013 7.18  Model Performance Figures", "49"),
    ("", ""),
    ("CHAPTER 8: CONCLUSION AND FUTURE SCOPE", "52"),
    ("8.1    Conclusion", "52"),
    ("8.2    Future Scope", "53"),
    ("", ""),
    ("CHAPTER 9: SUSTAINABLE DEVELOPMENT GOALS", "55"),
    ("9.1    SDG 8: Decent Work and Economic Growth", "55"),
    ("9.2    SDG 9: Industry, Innovation and Infrastructure", "56"),
    ("9.3    SDG 16: Peace, Justice and Strong Institutions", "56"),
    ("9.4    Broader Impact", "57"),
    ("9.5    Future Contribution to SDGs", "57"),
    ("", ""),
    ("REFERENCES", "58"),
]

toc_table = doc.add_table(rows=len(toc_entries), cols=2)
toc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (title, page) in enumerate(toc_entries):
    is_chapter = title.startswith("CHAPTER") or title == "REFERENCES"
    set_cell_text(toc_table.cell(i, 0), title, bold=is_chapter, font_size=11)
    set_cell_text(toc_table.cell(i, 1), page, font_size=11, align=WD_ALIGN_PARAGRAPH.RIGHT, bold=is_chapter)
    toc_table.cell(i, 0).width = Inches(5.0)
    toc_table.cell(i, 1).width = Inches(1.0)
    for cell in [toc_table.cell(i, 0), toc_table.cell(i, 1)]:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '</w:tcBorders>')
        tcPr.append(tcBorders)

# ============================================================
# LIST OF FIGURES
# ============================================================
add_page_break()
add_centered_text("LIST OF FIGURES", font_size=16, bold=True, space_before=24, space_after=12)

figures_list = [
    ("Fig. 4.1", "System Architecture Diagram", "22"),
    ("Fig. 4.2", "Use Case Diagram", "23"),
    ("Fig. 4.3", "ML Pipeline Diagram", "24"),
    ("Fig. 4.4", "Data Preprocessing Pipeline", "25"),
    ("Fig. 4.5", "Activity Diagram", "26"),
    ("Fig. 5.1", "Agile Development Model", "31"),
    ("Fig. 7.1", "Login Page", "42"),
    ("Fig. 7.2", "Registration Page", "42"),
    ("Fig. 7.3", "Invalid Login Attempt", "43"),
    ("Fig. 7.4", "Duplicate Registration Attempt", "43"),
    ("Fig. 7.5", "Home Page Dashboard", "44"),
    ("Fig. 7.6", "Prediction Form (Empty)", "44"),
    ("Fig. 7.7", "Prediction Form (Filled)", "45"),
    ("Fig. 7.8", "Prediction Result", "45"),
    ("Fig. 7.9", "Prediction History", "46"),
    ("Fig. 7.10", "EDA Visualization Gallery", "46"),
    ("Fig. 7.11", "Model Dashboard", "47"),
    ("Fig. 7.12", "Dashboard Charts (Detailed)", "47"),
    ("Fig. 7.13", "About Page", "48"),
    ("Fig. 7.14", "Model Accuracy Comparison", "49"),
    ("Fig. 7.15", "Confusion Matrix (Gradient Boosting)", "49"),
    ("Fig. 7.16", "Feature Importance", "50"),
    ("Fig. 7.17", "System Architecture", "50"),
    ("Fig. 7.18", "Data Preprocessing Pipeline", "51"),
]

lof_table = doc.add_table(rows=len(figures_list)+1, cols=3)
lof_table.style = 'Table Grid'
lof_table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_cell_text(lof_table.cell(0, 0), "Fig. No.", bold=True, font_size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_text(lof_table.cell(0, 1), "Title", bold=True, font_size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_text(lof_table.cell(0, 2), "Page No.", bold=True, font_size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
shade_cell(lof_table.cell(0, 0))
shade_cell(lof_table.cell(0, 1))
shade_cell(lof_table.cell(0, 2))
for i, (fno, ftitle, fpage) in enumerate(figures_list):
    r = i + 1
    set_cell_text(lof_table.cell(r, 0), fno, font_size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(lof_table.cell(r, 1), ftitle, font_size=10)
    set_cell_text(lof_table.cell(r, 2), fpage, font_size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    lof_table.cell(r, 0).width = Inches(0.8)
    lof_table.cell(r, 1).width = Inches(4.5)
    lof_table.cell(r, 2).width = Inches(0.9)

# ============================================================
# LIST OF TABLES
# ============================================================
add_page_break()
add_centered_text("LIST OF TABLES", font_size=16, bold=True, space_before=24, space_after=12)

tables_list = [
    ("Table 1.1", "Comparison of Existing Systems", "5"),
    ("Table 2.1", "Literature Survey Summary", "14"),
    ("Table 3.1", "Feasibility Study", "15"),
    ("Table 3.2", "Overall Description", "16"),
    ("Table 3.3", "System Feature Requirements", "17"),
    ("Table 3.4", "Non-Functional Requirements", "18"),
    ("Table 3.5", "Hardware Requirements", "19"),
    ("Table 3.6", "Software Requirements", "20"),
    ("Table 4.1", "Database Schema \u2013 Users Table", "27"),
    ("Table 4.2", "Database Schema \u2013 Predictions Table", "28"),
    ("Table 4.3", "Dataset Feature Descriptions", "28"),
    ("Table 6.1", "Test Cases \u2013 Authentication", "39"),
    ("Table 6.2", "Test Cases \u2013 Prediction", "40"),
    ("Table 6.3", "Test Cases \u2013 Visualization & Dashboard", "40"),
    ("Table 6.4", "Test Cases \u2013 Data & Security", "41"),
    ("Table 7.1", "Model Performance Comparison", "47"),
]

lot_table = doc.add_table(rows=len(tables_list)+1, cols=3)
lot_table.style = 'Table Grid'
lot_table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_cell_text(lot_table.cell(0, 0), "Table No.", bold=True, font_size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_text(lot_table.cell(0, 1), "Title", bold=True, font_size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_text(lot_table.cell(0, 2), "Page No.", bold=True, font_size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
shade_cell(lot_table.cell(0, 0))
shade_cell(lot_table.cell(0, 1))
shade_cell(lot_table.cell(0, 2))
for i, (tno, ttitle, tpage) in enumerate(tables_list):
    r = i + 1
    set_cell_text(lot_table.cell(r, 0), tno, font_size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(lot_table.cell(r, 1), ttitle, font_size=10)
    set_cell_text(lot_table.cell(r, 2), tpage, font_size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    lot_table.cell(r, 0).width = Inches(0.9)
    lot_table.cell(r, 1).width = Inches(4.4)
    lot_table.cell(r, 2).width = Inches(0.9)

# ============================================================
# SWITCH TO ARABIC PAGE NUMBERING
# ============================================================
new_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
add_page_number(new_section, start=1, fmt='decimal')

# ============================================================
# CHAPTER 1 - INTRODUCTION
# ============================================================
p_ch1 = add_centered_text("CHAPTER 1", font_size=18, bold=True, space_before=24, space_after=3)
p_ch1.paragraph_format.keep_with_next = True
p_ch1t = add_centered_text("INTRODUCTION", font_size=16, bold=True, space_after=10)
p_ch1t.paragraph_format.keep_with_next = True

add_section_heading("1.1", "Introduction")

add_justified_text(
    "Car insurance is one of the most essential financial products in the modern economy, providing "
    "protection to vehicle owners against financial losses arising from accidents, theft, natural disasters, "
    "and liability claims. The global car insurance market generates hundreds of billions of dollars in "
    "premiums annually, and the ability to accurately predict which policyholders are likely to file claims "
    "is fundamental to the profitability and sustainability of insurance companies. Accurate claim prediction "
    "enables insurers to set appropriate premium rates, maintain adequate reserves, and identify high-risk "
    "policyholders who may require additional underwriting scrutiny.",
    first_line_indent=1.27
)

add_justified_text(
    "Understanding the factors that influence insurance claim likelihood is crucial for risk assessment, "
    "premium pricing, and fraud detection. Traditional actuarial approaches rely on statistical models such "
    "as logistic regression and generalized linear models (GLMs) to estimate claim probabilities based on "
    "policyholder demographics and driving history. However, these methods assume linear relationships and "
    "may fail to capture the complex, non-linear interactions between variables such as driving experience, "
    "credit score, vehicle type, annual mileage, and past accident history that collectively determine "
    "claim probability.",
    first_line_indent=1.27
)

add_justified_text(
    "Machine learning offers a powerful alternative to traditional actuarial methods by automatically "
    "discovering complex patterns and non-linear relationships in policyholder data without requiring "
    "explicit mathematical formulations. Classification algorithms such as Random Forest, Gradient Boosting, "
    "and Support Vector Machines can model intricate feature interactions, handle categorical variables, and "
    "provide feature importance rankings that offer interpretable insights into which factors most strongly "
    "influence claim prediction outcomes.",
    first_line_indent=1.27
)

add_justified_text(
    "This project presents a comprehensive web-based platform for car insurance claim prediction using four "
    "machine learning classification models: Random Forest, Gradient Boosting, Support Vector Machine (SVM), "
    "and Logistic Regression. The system processes a synthetic Car Insurance Claim dataset containing 10,000 "
    "records across 17 features, performs automated data preprocessing (label encoding, standard scaling), "
    "trains all four models with an 80/20 stratified train-test split, and presents results through "
    "Chart.js visualizations and an interactive dashboard.",
    first_line_indent=1.27
)

add_justified_text(
    "The Flask web framework was chosen for its simplicity, flexibility, and seamless integration with "
    "Python's scientific computing ecosystem. The application uses pandas for data manipulation, scikit-learn "
    "for model training and evaluation, Chart.js for interactive chart generation, and SQLite for persistent "
    "user authentication and prediction history storage. The Bootstrap 5 dark-themed user interface provides "
    "a modern, professional appearance with responsive design, ensuring an engaging user experience for "
    "insurance professionals, data analysts, and students.",
    first_line_indent=1.27
)

add_section_heading("1.2", "Scope of the Project")

add_justified_text(
    "The scope of this project encompasses the design, development, and evaluation of an end-to-end machine "
    "learning pipeline for car insurance claim prediction. The key areas covered include:",
    first_line_indent=1.27
)
add_bullet("Implementation of four classification models (Random Forest, Gradient Boosting, SVM, Logistic Regression) for comprehensive performance comparison on the insurance claim dataset.")
add_bullet("Automated data preprocessing pipeline including label encoding of 9 categorical features, standard scaling of 8 numeric features, and stratified 80/20 train-test split preserving class distribution.")
add_bullet("Interactive web interface with user authentication (registration, login, logout), prediction form with 17 input fields, and real-time claim prediction with confidence scores.")
add_bullet("Exploratory data analysis through Chart.js visualizations including distribution charts, correlation analyses, feature importance plots, and model comparison bar charts.")
add_bullet("Prediction history tracking with SQLite database storage, enabling users to review past predictions and analyze trends in their assessments.")
add_bullet("Comprehensive dashboard with model analytics showing accuracy, precision, recall, and F1-score metrics for all four classification models.")
add_bullet("Docker containerization for streamlined deployment and reproducibility across different computing environments.")
add_bullet("Secure authentication system using Werkzeug password hashing to protect user accounts and prediction data.")

add_section_heading("1.3", "Objectives")

add_justified_text(
    "The primary objectives of this project are:",
    first_line_indent=1.27
)
add_bullet("To design and implement a machine learning pipeline capable of accurately predicting car insurance claim outcomes from policyholder demographics and driving history.")
add_bullet("To compare the performance of four classification algorithms (Random Forest, Gradient Boosting, SVM, Logistic Regression) using accuracy, precision, recall, and F1-score metrics.")
add_bullet("To develop a user-friendly web application with authentication for inputting policyholder details and receiving automated claim predictions with confidence scores.")
add_bullet("To implement automated data preprocessing including label encoding for categorical variables and standard scaling for numeric features.")
add_bullet("To create interactive Chart.js visualizations for exploratory data analysis including distribution charts, feature importance plots, and model comparison charts.")
add_bullet("To implement a prediction history system using SQLite database for tracking and reviewing past predictions.")
add_bullet("To demonstrate the practical applicability of machine learning classification techniques for insurance risk assessment.")
add_bullet("To containerize the application using Docker for consistent deployment across different environments.")

add_section_heading("1.4", "Problem Formulation")

add_justified_text(
    "Predicting car insurance claims accurately is a complex challenge due to the multifaceted nature of the "
    "factors involved. The current approaches to insurance claim analysis face several key challenges:",
    first_line_indent=1.27
)

add_justified_text(
    "First, the insurance claim dataset exhibits significant class imbalance, with approximately 74% of "
    "policyholders not filing claims (OUTCOME=0) and only 26% filing claims (OUTCOME=1). This imbalance "
    "means that a naive classifier predicting 'no claim' for all instances would achieve 74% accuracy, "
    "making it crucial to evaluate models using precision, recall, and F1-score in addition to accuracy. "
    "The stratified train-test split preserves this class distribution in both training and testing sets.",
    first_line_indent=1.27
)

add_justified_text(
    "Second, the relationship between policyholder features and claim likelihood is inherently non-linear. "
    "For example, the effect of driving experience on claim probability follows a non-linear pattern where "
    "very new drivers and very experienced drivers may have different risk profiles. Similarly, the interaction "
    "between credit score and annual mileage may create complex risk patterns that linear models cannot capture "
    "without extensive feature engineering, while ensemble methods like Gradient Boosting can model them "
    "automatically.",
    first_line_indent=1.27
)

add_justified_text(
    "Third, the dataset contains a mix of categorical and numeric features that require different preprocessing "
    "approaches. The 9 categorical features (gender, race, driving experience, education, income, vehicle "
    "ownership, vehicle year, married, vehicle type) need label encoding to convert them to numeric values, "
    "while the 8 numeric features (age, credit score, annual mileage, speeding violations, DUIs, past "
    "accidents, postal code, children) benefit from standard scaling to normalize their ranges for "
    "distance-based algorithms like SVM.",
    first_line_indent=1.27
)

add_justified_text(
    "Fourth, existing tools for insurance claim analysis are primarily limited to proprietary actuarial "
    "software and statistical packages that require specialized expertise. There is a significant gap between "
    "the analytical capabilities of machine learning and the accessibility of these tools for insurance "
    "professionals who may not have technical programming backgrounds. A web-based interface with "
    "automated preprocessing and model selection would democratize access to advanced predictive analytics "
    "in the insurance domain.",
    first_line_indent=1.27
)

add_section_heading("1.5", "Existing System")

add_justified_text(
    "The current approaches to car insurance claim prediction can be broadly categorized into manual "
    "actuarial analysis using traditional software, statistical modeling with specialized tools, and "
    "automated machine learning (AutoML) platforms. Manual actuarial analysis using tools like SAS, SPSS, "
    "or R requires significant domain expertise and programming knowledge, limiting accessibility to "
    "trained actuaries and data scientists.",
    first_line_indent=1.27
)

add_justified_text(
    "Statistical modeling tools provide a flexible environment for building claim prediction models but "
    "lack user-friendly interfaces for non-technical insurance professionals. While these tools excel at "
    "implementing GLMs and survival analysis models, they require programming knowledge, do not provide "
    "built-in web deployment capabilities, and cannot generate interactive dashboards automatically. "
    "AutoML platforms like DataRobot and H2O.ai offer automated model selection but are typically "
    "cloud-based, involve significant costs, and may not be suitable for sensitive policyholder data.",
    first_line_indent=1.27
)

add_centered_text("Table 1.1: Comparison of Existing Systems", font_size=10, bold=True, space_after=4, keep_with_next=True)
exist_table = doc.add_table(rows=5, cols=4)
exist_table.style = 'Table Grid'
exist_table.alignment = WD_TABLE_ALIGNMENT.CENTER
exist_data = [
    ("System", "Method", "Strengths", "Limitations"),
    ("Manual Actuarial\nAnalysis (SAS/SPSS)", "GLMs, logistic\nregression, frequency-\nseverity models", "Well-established\nmethods, interpretable", "Requires expertise, limited\nto linear models, slow"),
    ("Statistical Tools\n(R/Python notebooks)", "Flexible ML pipeline,\ncustom visualizations", "Highly customizable,\nreproducible", "Requires programming,\nno web interface, no auth"),
    ("AutoML Platforms\n(DataRobot, H2O.ai)", "Automated model\nselection and tuning", "Minimal coding,\nbroad model search", "Cloud-dependent, costly,\nprivacy concerns"),
    ("Our Proposed System", "4 ML classification\nmodels, web UI, auth", "No coding required,\ninteractive, local, secure", "Limited to structured\ninput, single dataset format"),
]
for i, row_data in enumerate(exist_data):
    for j, val in enumerate(row_data):
        set_cell_text(exist_table.cell(i, j), val, bold=(i == 0), font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER if j != 3 else WD_ALIGN_PARAGRAPH.LEFT)
    if i == 0:
        for j in range(4):
            shade_cell(exist_table.cell(i, j))
    exist_table.cell(i, 0).width = Inches(1.3)
    exist_table.cell(i, 1).width = Inches(1.5)
    exist_table.cell(i, 2).width = Inches(1.4)
    exist_table.cell(i, 3).width = Inches(2.0)
keep_table_on_one_page(exist_table)

add_justified_text(
    "The disadvantages of existing systems include:",
    space_before=6
)
add_bullet("High dependency on programming expertise for data analysis and model development.")
add_bullet("Lack of user-friendly web interfaces with authentication for insurance professionals.")
add_bullet("Limited prediction history tracking and dashboard capabilities for ongoing risk assessment.")
add_bullet("Cloud-based solutions raise data privacy and cost concerns for sensitive policyholder data.")
add_bullet("No integrated platform that combines preprocessing, multiple model comparison, interactive visualization, authentication, and prediction history.")

add_section_heading("1.6", "Proposed System")

add_justified_text(
    "The proposed system addresses the limitations of existing approaches by providing a self-contained, "
    "web-based machine learning platform for car insurance claim prediction. The system includes user "
    "authentication with secure password hashing, prediction history tracking with SQLite storage, "
    "and comprehensive model analytics through an interactive dashboard. Users register, log in, input "
    "policyholder details through a structured form, and receive claim predictions with confidence scores.",
    first_line_indent=1.27
)

add_justified_text(
    "The key advantages of the proposed system include:",
    space_before=4
)
add_bullet("Four classification models (Random Forest, Gradient Boosting, SVM, Logistic Regression) with the best model (Gradient Boosting, 91.95% accuracy) used for predictions.")
add_bullet("Automated data preprocessing pipeline that handles label encoding of categorical features and standard scaling of numeric features without user intervention.")
add_bullet("Secure user authentication using Werkzeug password hashing with SQLite database storage for user accounts and prediction history.")
add_bullet("Interactive Chart.js visualizations including distribution charts, feature importance plots, model comparison charts, and confusion matrices.")
add_bullet("Prediction history page allowing users to review past predictions and track assessment patterns over time.")
add_bullet("Comprehensive dashboard with model analytics showing accuracy, precision, recall, and F1-score for all four models.")
add_bullet("Bootstrap 5 dark-themed responsive user interface with modern design and intuitive navigation.")
add_bullet("Docker containerization for one-command deployment on any machine with Docker installed.")
add_bullet("Completely local processing with no data sent to external servers, ensuring policyholder data privacy.")

add_justified_text(
    "The proposed system follows a Model-View-Controller (MVC) architecture with Flask handling HTTP "
    "requests, SQLite providing persistent storage for users and predictions, scikit-learn powering the "
    "ML classification pipeline, and Chart.js generating interactive visualizations. The pre-trained "
    "Gradient Boosting model is loaded at application startup for fast prediction responses, while the "
    "dashboard provides comprehensive analytics comparing all four trained models.",
    first_line_indent=1.27
)

# ============================================================
# CHAPTER 2 - LITERATURE SURVEY
# ============================================================
p_ch2 = add_centered_text("CHAPTER 2", font_size=18, bold=True, space_before=24, space_after=3)
p_ch2.paragraph_format.keep_with_next = True
p_ch2.paragraph_format.page_break_before = True
p_ch2t = add_centered_text("LITERATURE SURVEY", font_size=16, bold=True, space_after=10)
p_ch2t.paragraph_format.keep_with_next = True

# 2.1
add_section_heading("2.1", "Machine Learning for Insurance Claim Prediction")
add_justified_text(
    "Noll et al. (2018) demonstrated the potential of machine learning for insurance pricing and claim "
    "prediction. Their comprehensive study compared traditional generalized linear models (GLMs) with "
    "modern ML algorithms including Random Forest, Gradient Boosting, and neural networks for predicting "
    "claim frequency and severity. The authors found that ensemble methods consistently outperformed "
    "traditional actuarial models, particularly when applied to datasets with non-linear relationships "
    "between policyholder features and claim outcomes.",
    first_line_indent=1.27
)
add_justified_text(
    "The study identified several key challenges in insurance claim prediction: class imbalance (most "
    "policyholders do not file claims), heterogeneous feature types (categorical and numeric), and the "
    "need for interpretable models in regulatory contexts. For car insurance specifically, the authors "
    "noted that driving experience, vehicle characteristics, and geographic location interact in complex "
    "ways that GLMs cannot adequately capture without extensive feature engineering.",
    first_line_indent=1.27
)

# 2.2
add_section_heading("2.2", "Random Forest for Classification Tasks")
add_justified_text(
    "Breiman (2001) introduced the Random Forest algorithm, an ensemble learning method that constructs "
    "multiple decision trees during training and outputs the majority vote of their individual predictions "
    "for classification tasks. Random Forest addresses the overfitting problem inherent in single decision "
    "trees by introducing randomness in both the data sampling (bootstrap aggregating or bagging) and "
    "feature selection at each split point. The algorithm has become one of the most widely used "
    "classification methods due to its robustness and built-in feature importance estimation.",
    first_line_indent=1.27
)
add_justified_text(
    "For insurance data classification, Random Forest offers particular advantages: it handles mixed "
    "feature types gracefully, is robust to outliers, does not require feature scaling, and provides "
    "out-of-bag error estimates for model validation. The feature importance scores generated by Random "
    "Forest reveal which policyholder characteristics most strongly influence claim predictions, providing "
    "actionable insights for underwriting. In our project, Random Forest achieved 90.10% accuracy with "
    "86.10% precision and 73.85% recall.",
    first_line_indent=1.27
)

# 2.3
add_section_heading("2.3", "Gradient Boosting Methods")
add_justified_text(
    "Friedman (2001) developed the Gradient Boosting Machine (GBM), a sequential ensemble method that "
    "builds decision trees iteratively, with each new tree correcting the errors of the previous ensemble. "
    "Unlike Random Forest which builds trees independently in parallel, Gradient Boosting constructs trees "
    "sequentially, fitting each new tree to the negative gradient of the loss function. This sequential "
    "optimization approach often achieves higher accuracy than Random Forest at the cost of longer training "
    "times and greater sensitivity to hyperparameters.",
    first_line_indent=1.27
)
add_justified_text(
    "Chen and Guestrin (2016) introduced XGBoost, an optimized implementation of gradient boosting that "
    "has dominated machine learning competitions. While our project uses scikit-learn's "
    "GradientBoostingClassifier, the underlying principles are the same: sequential tree construction with "
    "gradient-based optimization. Gradient Boosting achieved the best performance in our insurance claim "
    "prediction task with 91.95% accuracy, 89.45% precision, 78.27% recall, and 83.49% F1-score, "
    "demonstrating its ability to capture non-linear risk patterns.",
    first_line_indent=1.27
)

# 2.4
add_section_heading("2.4", "Support Vector Machines for Classification")
add_justified_text(
    "Cortes and Vapnik (1995) introduced the Support Vector Machine (SVM) algorithm for binary "
    "classification. SVM finds the optimal hyperplane that maximally separates two classes in feature "
    "space, with support vectors being the data points closest to the decision boundary. The kernel trick "
    "enables SVM to handle non-linearly separable data by mapping features into higher-dimensional spaces "
    "where linear separation is possible. Common kernels include linear, polynomial, and radial basis "
    "function (RBF).",
    first_line_indent=1.27
)
add_justified_text(
    "For insurance claim classification, SVM with RBF kernel can capture complex non-linear boundaries "
    "between claim and no-claim regions in the feature space. However, SVM requires feature scaling "
    "(standardization) for optimal performance, as features with larger ranges can dominate the distance "
    "calculations. In our project, SVM achieved 91.10% accuracy with 85.62% precision and 79.04% recall "
    "after standard scaling of all numeric features.",
    first_line_indent=1.27
)

# 2.5
add_section_heading("2.5", "Logistic Regression for Binary Classification")
add_justified_text(
    "Hosmer, Lemeshow, and Sturdivant (2013) provided a comprehensive treatment of logistic regression, "
    "the foundational algorithm for binary classification. Logistic regression models the probability of "
    "the positive class using the sigmoid function applied to a linear combination of features. Despite "
    "its simplicity, logistic regression provides valuable insights through its coefficient values, which "
    "indicate the direction and magnitude of each feature's effect on claim probability, and through its "
    "predicted probabilities, which serve as confidence scores.",
    first_line_indent=1.27
)
add_justified_text(
    "In our insurance claim prediction project, Logistic Regression serves as the baseline model against "
    "which more complex algorithms are compared. Its accuracy of 90.25% with 83.92% precision and 77.31% "
    "recall indicates that a linear combination of the policyholder features captures most of the claim "
    "pattern. The improvement achieved by non-linear models (Gradient Boosting: 91.95%) quantifies the "
    "extent of non-linear relationships in the data that logistic regression cannot capture.",
    first_line_indent=1.27
)

# 2.6
add_section_heading("2.6", "Feature Engineering for Insurance Data")
add_justified_text(
    "Kuhn and Johnson (2013) authored a comprehensive guide on feature engineering and selection for "
    "predictive modeling. They emphasized that the quality of input features often matters more than the "
    "choice of algorithm, and that domain knowledge should guide feature transformation decisions. For "
    "insurance datasets, common feature engineering techniques include encoding categorical variables "
    "(label encoding, one-hot encoding), scaling numeric features (standardization, min-max), and creating "
    "interaction terms between related policyholder attributes.",
    first_line_indent=1.27
)
add_justified_text(
    "In the Car Insurance Claim dataset, preprocessing decisions significantly impact model performance. "
    "Label encoding converts 9 categorical features (gender, race, driving experience, education, income, "
    "vehicle ownership, vehicle year, married, vehicle type) to numeric values suitable for all four "
    "classification models. Standard scaling normalizes the 8 numeric features (age, credit score, annual "
    "mileage, speeding violations, DUIs, past accidents, postal code, children) to zero mean and unit "
    "variance, which is critical for SVM performance and beneficial for logistic regression convergence.",
    first_line_indent=1.27
)

# 2.7
add_section_heading("2.7", "Class Imbalance in Insurance Datasets")
add_justified_text(
    "He and Garcia (2009) provided a comprehensive review of learning from imbalanced data, a common "
    "challenge in insurance claim prediction where the majority of policyholders do not file claims. "
    "They discussed techniques including oversampling (SMOTE), undersampling, cost-sensitive learning, "
    "and ensemble methods designed for imbalanced classification. The authors noted that standard accuracy "
    "is a misleading metric for imbalanced datasets, recommending precision, recall, F1-score, and "
    "AUC-ROC as more informative evaluation criteria.",
    first_line_indent=1.27
)
add_justified_text(
    "Our Car Insurance Claim dataset exhibits moderate class imbalance with approximately 74% no-claim "
    "(OUTCOME=0) and 26% claim (OUTCOME=1) instances. The stratified train-test split preserves this "
    "distribution in both training and testing sets. We evaluate models using accuracy, precision, recall, "
    "and F1-score to provide a comprehensive view of model performance on both classes. The Gradient "
    "Boosting model's balanced performance (89.45% precision, 78.27% recall) indicates effective handling "
    "of the class imbalance without explicit resampling techniques.",
    first_line_indent=1.27
)

# 2.8
add_section_heading("2.8", "Synthetic Data Generation for Insurance Research")
add_justified_text(
    "Patki et al. (2016) introduced techniques for generating synthetic tabular data that preserves the "
    "statistical properties of real datasets while protecting privacy. Synthetic data generation is "
    "particularly valuable in insurance research where real policyholder data is confidential and subject "
    "to strict privacy regulations (GDPR, HIPAA). The authors demonstrated that models trained on "
    "well-crafted synthetic data can achieve comparable performance to those trained on real data.",
    first_line_indent=1.27
)
add_justified_text(
    "Our project uses a synthetic Car Insurance Claim dataset containing 10,000 records generated using "
    "the generate_dataset.py script. The synthetic data preserves realistic distributions for all 17 "
    "features: ages follow a normal distribution (18-70), credit scores are uniformly distributed "
    "(300-850), and the claim outcome (OUTCOME) is correlated with risk factors like speeding violations, "
    "DUIs, and past accidents. This approach enables reproducible research without privacy concerns while "
    "maintaining the statistical properties needed for meaningful ML model evaluation.",
    first_line_indent=1.27
)

# 2.9
add_section_heading("2.9", "Flask for Data Science Web Applications")
add_justified_text(
    "Grinberg (2018) authored the definitive guide on Flask web development, documenting the framework's "
    "design philosophy of minimalism and extensibility. Flask's microframework architecture provides core "
    "HTTP request handling, URL routing, and template rendering while allowing developers to choose their "
    "preferred libraries for data processing, authentication, and visualization. This modularity makes "
    "Flask particularly suitable for data science applications where rapid integration with Python "
    "scientific computing libraries is a priority.",
    first_line_indent=1.27
)
add_justified_text(
    "For our car insurance claim prediction platform, Flask provides the web layer connecting the user "
    "interface to the ML pipeline. The framework's support for Jinja2 templates enables creation of "
    "dynamic pages that embed Chart.js visualizations, display prediction results with confidence scores, "
    "and provide navigation between authentication, prediction, history, visualization, and dashboard "
    "pages. Flask's session management handles user authentication state across requests.",
    first_line_indent=1.27
)

# 2.10
add_section_heading("2.10", "Chart.js for Interactive Visualization")
add_justified_text(
    "Chart.js (2023) is an open-source JavaScript charting library that provides simple yet flexible "
    "chart creation capabilities for web applications. Unlike server-side visualization libraries like "
    "Matplotlib, Chart.js renders charts directly in the browser using HTML5 Canvas, enabling smooth "
    "animations, responsive sizing, and interactive features like tooltips, legends, and click events. "
    "Chart.js supports bar charts, line charts, doughnut charts, radar charts, and scatter plots out of "
    "the box.",
    first_line_indent=1.27
)
add_justified_text(
    "In our insurance claim prediction platform, Chart.js generates multiple interactive visualizations: "
    "distribution charts showing feature value distributions across the dataset, model comparison bar "
    "charts displaying accuracy, precision, recall, and F1-score for all four models, confusion matrix "
    "heatmaps for the best model, and feature importance horizontal bar charts. These client-side "
    "charts provide smooth interactivity without requiring server round-trips.",
    first_line_indent=1.27
)

# 2.11
add_section_heading("2.11", "SQLite for Lightweight Database Storage")
add_justified_text(
    "Owens (2006) described SQLite as a self-contained, serverless, zero-configuration relational database "
    "engine that is ideal for embedded applications and small-to-medium web applications. Unlike client-server "
    "databases (MySQL, PostgreSQL), SQLite stores the entire database in a single file, requiring no separate "
    "server process, no configuration, and no administration. This makes SQLite perfect for applications "
    "that need persistent storage without the complexity of a full database management system.",
    first_line_indent=1.27
)
add_justified_text(
    "Our car insurance claim prediction platform uses SQLite (insurance.db) with two tables: users (storing "
    "usernames, hashed passwords, names, and roles) and predictions (storing input data as JSON, prediction "
    "outcomes, confidence scores, and timestamps). SQLite's file-based architecture simplifies deployment "
    "(no database server to configure), Docker containerization (single file to persist), and backup "
    "(copy one file). The database integrates seamlessly with Flask's request lifecycle.",
    first_line_indent=1.27
)

# 2.12
add_section_heading("2.12", "Model Evaluation Metrics for Classification")
add_justified_text(
    "Powers (2011) provided a comprehensive analysis of evaluation metrics for binary classification, "
    "distinguishing between accuracy, precision, recall, F1-score, and AUC-ROC. Accuracy measures the "
    "proportion of correct predictions overall, but can be misleading with imbalanced classes. Precision "
    "measures the proportion of predicted positives that are actually positive (relevant for minimizing "
    "false alarms). Recall measures the proportion of actual positives that are correctly identified "
    "(relevant for catching all true claims). F1-score is the harmonic mean of precision and recall.",
    first_line_indent=1.27
)
add_justified_text(
    "Our project evaluates all four models using accuracy, precision, recall, and F1-score, providing "
    "complementary views of classification quality. The Gradient Boosting model achieves the best balance "
    "with 91.95% accuracy, 89.45% precision (few false claim predictions), 78.27% recall (catches most "
    "actual claims), and 83.49% F1-score. This comprehensive evaluation enables informed model selection "
    "based on the specific requirements of the insurance use case.",
    first_line_indent=1.27
)

# 2.13
add_section_heading("2.13", "Web Authentication and Security")
add_justified_text(
    "OWASP (2021) published comprehensive guidelines for web application security, covering authentication, "
    "session management, input validation, and password storage. The guidelines recommend storing passwords "
    "as salted hashes using strong algorithms (bcrypt, scrypt, Argon2, or PBKDF2) rather than plain text "
    "or simple hashes. Session tokens should be generated with sufficient entropy and transmitted securely. "
    "Input validation should be applied to all user inputs to prevent injection attacks.",
    first_line_indent=1.27
)
add_justified_text(
    "Our platform implements secure authentication using Werkzeug's generate_password_hash() and "
    "check_password_hash() functions, which use PBKDF2 with SHA-256 by default. User passwords are never "
    "stored in plain text. Flask's session management provides secure, signed cookies for maintaining "
    "authentication state. The @login_required decorator pattern protects all sensitive routes (predict, "
    "history, dashboard, visualize) from unauthorized access.",
    first_line_indent=1.27
)

# 2.14
add_section_heading("2.14", "Bootstrap 5 for Responsive Web Design")
add_justified_text(
    "Bootstrap (2021) released version 5 of the popular CSS framework, removing the jQuery dependency and "
    "introducing improved grid system, utility classes, and component updates. Bootstrap 5's responsive "
    "design system uses a mobile-first approach with breakpoints for different screen sizes, ensuring that "
    "web applications render correctly on devices ranging from smartphones to large desktop monitors. The "
    "dark theme support in Bootstrap 5 enables creation of modern, eye-friendly interfaces.",
    first_line_indent=1.27
)
add_justified_text(
    "Our car insurance claim prediction platform uses Bootstrap 5 with a dark theme for the entire "
    "application. The responsive grid system ensures that prediction forms, history tables, visualization "
    "galleries, and dashboard analytics display correctly across all device sizes. Bootstrap's card "
    "components, navigation bars, and form controls provide a consistent, professional appearance while "
    "minimizing custom CSS requirements.",
    first_line_indent=1.27
)

# 2.15
add_section_heading("2.15", "Docker for ML Application Deployment")
add_justified_text(
    "Merkel (2014) described Docker as a platform for creating, deploying, and running applications in "
    "lightweight, portable containers. Docker containers package an application with all its dependencies "
    "into a standardized unit that runs consistently across different computing environments. This eliminates "
    "the \u201cworks on my machine\u201d problem by ensuring identical development, testing, and production "
    "environments. Docker's layered filesystem and image caching mechanisms minimize storage overhead.",
    first_line_indent=1.27
)
add_justified_text(
    "Our car insurance claim prediction system includes a Dockerfile that builds a complete application "
    "image from the Python 3.11-slim base image. The Docker build process installs all Python dependencies "
    "(Flask, pandas, scikit-learn, werkzeug) and configures the Flask application to run on port 5002. "
    "This containerization approach enables one-command deployment (docker build && docker run) and ensures "
    "that the system can be reproduced on any machine with Docker installed, regardless of the host "
    "operating system or Python version.",
    first_line_indent=1.27
)

# Literature Survey Summary Table
add_page_break()
add_centered_text("Table 2.1: Literature Survey Summary", font_size=10, bold=True, space_before=6, space_after=4, keep_with_next=True)
lit_table = doc.add_table(rows=16, cols=4)
lit_table.style = 'Table Grid'
lit_table.alignment = WD_TABLE_ALIGNMENT.CENTER

lit_data = [
    ("S.No.", "Author(s) / Year", "Topic", "Key Contribution"),
    ("1", "Noll et al. (2018)", "ML for Insurance", "Ensemble methods outperform GLMs for claim prediction"),
    ("2", "Breiman (2001)", "Random Forest", "Bagging + random features; robust classification ensemble"),
    ("3", "Friedman (2001)", "Gradient Boosting", "Sequential tree optimization; high accuracy classification"),
    ("4", "Cortes & Vapnik (1995)", "SVM", "Maximum margin classifier; kernel trick for non-linear data"),
    ("5", "Hosmer et al. (2013)", "Logistic Regression", "Sigmoid function; baseline binary classification"),
    ("6", "Kuhn & Johnson (2013)", "Feature Engineering", "Feature quality vs. algorithm choice for prediction"),
    ("7", "He & Garcia (2009)", "Class Imbalance", "Evaluation metrics and resampling for imbalanced data"),
    ("8", "Patki et al. (2016)", "Synthetic Data", "Privacy-preserving synthetic data for ML research"),
    ("9", "Grinberg (2018)", "Flask Framework", "Lightweight web development for ML integration"),
    ("10", "Chart.js (2023)", "Chart.js", "Client-side interactive HTML5 Canvas visualizations"),
    ("11", "Owens (2006)", "SQLite", "Serverless embedded relational database engine"),
    ("12", "Powers (2011)", "Classification Metrics", "Precision, recall, F1-score for binary classification"),
    ("13", "OWASP (2021)", "Web Security", "Authentication, password hashing, session management"),
    ("14", "Bootstrap (2021)", "Bootstrap 5", "Responsive dark-theme CSS framework"),
    ("15", "Merkel (2014)", "Docker", "Containerization; reproducible deployment"),
]

for i, row_data in enumerate(lit_data):
    for j, val in enumerate(row_data):
        set_cell_text(lit_table.cell(i, j), val, bold=(i == 0), font_size=8, align=WD_ALIGN_PARAGRAPH.CENTER if j == 0 else WD_ALIGN_PARAGRAPH.LEFT)
    if i == 0:
        for j in range(4):
            shade_cell(lit_table.cell(i, j))
    lit_table.cell(i, 0).width = Inches(0.4)
    lit_table.cell(i, 1).width = Inches(1.6)
    lit_table.cell(i, 2).width = Inches(1.4)
    lit_table.cell(i, 3).width = Inches(2.8)
keep_table_on_one_page(lit_table)

add_justified_text(
    "The literature survey demonstrates that machine learning classification methods, particularly ensemble "
    "approaches like Gradient Boosting and Random Forest, are highly effective for predicting insurance "
    "claim outcomes from policyholder data. The survey also highlights the importance of complementary "
    "technologies \u2014 Flask for web development, Chart.js for interactive visualization, SQLite for "
    "lightweight storage, Bootstrap 5 for responsive design, and Docker for deployment \u2014 in building "
    "practical, accessible insurance analytics platforms.",
    space_before=8
)

# ============================================================
# CHAPTER 3 - REQUIREMENT ANALYSIS
# ============================================================
p_ch3 = add_centered_text("CHAPTER 3", font_size=18, bold=True, space_before=24, space_after=3)
p_ch3.paragraph_format.keep_with_next = True
p_ch3.paragraph_format.page_break_before = True
p_ch3t = add_centered_text("REQUIREMENT ANALYSIS AND SYSTEM SPECIFICATION", font_size=16, bold=True, space_after=10)
p_ch3t.paragraph_format.keep_with_next = True

add_section_heading("3.1", "Feasibility Study")

add_justified_text(
    "A feasibility study evaluates the practicality and viability of a proposed project before significant "
    "resources are committed to its development. For this car insurance claim prediction system, the feasibility "
    "analysis covers three critical dimensions: technical feasibility, economic feasibility, and operational feasibility.",
    space_after=2, keep_with_next=True
)

add_centered_text("Table 3.1: Feasibility Study Table", font_size=10, bold=True, space_after=2, keep_with_next=True)
feas_table = doc.add_table(rows=4, cols=2)
feas_table.style = 'Table Grid'
feas_table.alignment = WD_TABLE_ALIGNMENT.CENTER
feas_data = [
    ("Feasibility Type", "Description"),
    ("Technical Feasibility", "The system uses Python, Flask, pandas, scikit-learn, Chart.js, Bootstrap 5, and SQLite \u2014 all mature, well-documented technologies with active communities. The four classification models require moderate computational resources for training (under 60 seconds on the 10,000-row dataset). Docker provides cross-platform deployment compatibility."),
    ("Economic Feasibility", "All technologies used are free and open-source (Python, Flask, scikit-learn, Chart.js, Bootstrap, SQLite). No licensing fees are required. The system runs on commodity hardware with 4GB+ RAM. The synthetic dataset is generated locally at no cost. The estimated development time is 6-8 weeks for a team of students."),
    ("Operational Feasibility", "The system provides an intuitive web interface with user authentication, structured prediction forms, and interactive dashboards. Users register once, log in, input policyholder details through a form, and receive immediate predictions with confidence scores. Prediction history is stored for future reference. The Flask application can be deployed via Docker for containerized deployment."),
]
for i, (ftype, desc) in enumerate(feas_data):
    set_cell_text(feas_table.cell(i, 0), ftype, bold=True, font_size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(feas_table.cell(i, 1), desc, font_size=10)
    feas_table.cell(i, 0).width = Inches(1.4)
    feas_table.cell(i, 1).width = Inches(4.8)
    if i == 0:
        shade_cell(feas_table.cell(i, 0))
        shade_cell(feas_table.cell(i, 1))
keep_table_on_one_page(feas_table)

add_section_heading("3.2", "Software Requirement Specification")

add_subsection_heading("3.2.1", "Overall Description")
add_centered_text("Table 3.2: Overall Description", font_size=10, bold=True, space_after=4, keep_with_next=True)
od_table = doc.add_table(rows=8, cols=2)
od_table.style = 'Table Grid'
od_table.alignment = WD_TABLE_ALIGNMENT.CENTER
od_data = [
    ("Parameter", "Description"),
    ("Product Name", "Car Insurance Claim Prediction Platform"),
    ("Product Type", "Web-based Machine Learning Classification Application"),
    ("Purpose", "Automated prediction of car insurance claim outcomes using 4 classification models"),
    ("Users", "Insurance professionals, underwriters, data analysts, students"),
    ("Platform", "Cross-platform (accessible via web browser)"),
    ("Database", "SQLite (insurance.db) for user accounts and prediction history"),
    ("Authentication", "Werkzeug password hashing with Flask session management"),
]
for i, (param, desc) in enumerate(od_data):
    set_cell_text(od_table.cell(i, 0), param, bold=(i == 0), font_size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(od_table.cell(i, 1), desc, bold=(i == 0), font_size=10)
    od_table.cell(i, 0).width = Inches(1.5)
    od_table.cell(i, 1).width = Inches(4.7)
    if i == 0:
        shade_cell(od_table.cell(i, 0))
        shade_cell(od_table.cell(i, 1))
keep_table_on_one_page(od_table)

p_sfr = add_subsection_heading("3.2.2", "System Feature Requirement")
p_sfr.paragraph_format.page_break_before = True
add_centered_text("Table 3.3: System Feature Requirements", font_size=10, bold=True, space_after=4, keep_with_next=True)
sfr_table = doc.add_table(rows=8, cols=2)
sfr_table.style = 'Table Grid'
sfr_table.alignment = WD_TABLE_ALIGNMENT.CENTER
sfr_data = [
    ("Feature", "Description"),
    ("User Authentication", "Secure registration and login with Werkzeug password hashing; role-based access control"),
    ("Claim Prediction", "17-field input form for policyholder details; real-time prediction with confidence score using pre-trained Gradient Boosting model"),
    ("Prediction History", "SQLite-stored history of past predictions with input data, outcomes, confidence scores, and timestamps"),
    ("EDA Visualizations", "Chart.js gallery showing feature distributions, correlation analyses, and data insights across the insurance dataset"),
    ("Model Dashboard", "Comprehensive analytics comparing accuracy, precision, recall, and F1-score for all 4 classification models"),
    ("Data Preprocessing", "Automated pipeline: label encoding for 9 categorical features, standard scaling for 8 numeric features, stratified 80/20 split"),
    ("About Page", "Project information, technology stack description, and team details"),
]
for i, (feat, desc) in enumerate(sfr_data):
    set_cell_text(sfr_table.cell(i, 0), feat, bold=(i == 0), font_size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(sfr_table.cell(i, 1), desc, bold=(i == 0), font_size=10)
    sfr_table.cell(i, 0).width = Inches(1.5)
    sfr_table.cell(i, 1).width = Inches(4.7)
    if i == 0:
        shade_cell(sfr_table.cell(i, 0))
        shade_cell(sfr_table.cell(i, 1))
keep_table_on_one_page(sfr_table)

p_nfr = add_subsection_heading("3.2.3", "Non-Functional Requirement")
add_centered_text("Table 3.4: Non-Functional Requirements", font_size=10, bold=True, space_after=4, keep_with_next=True)
nfr_table = doc.add_table(rows=7, cols=2)
nfr_table.style = 'Table Grid'
nfr_table.alignment = WD_TABLE_ALIGNMENT.CENTER
nfr_data = [
    ("Requirement", "Description"),
    ("Performance", "Prediction response under 1 second (pre-trained model); page load under 2 seconds; chart rendering under 3 seconds"),
    ("Usability", "Intuitive form-based interface; Bootstrap 5 dark theme; responsive design; clear prediction display with confidence score"),
    ("Scalability", "Modular architecture supports additional ML models, evaluation metrics, and visualization types"),
    ("Reliability", "Graceful error handling for invalid inputs, authentication failures, and malformed data; SQLite transaction safety"),
    ("Portability", "Docker containerization; cross-platform Flask deployment; SQLite requires no database server"),
    ("Security", "Werkzeug password hashing (PBKDF2-SHA256); Flask session management; @login_required route protection; no plain-text passwords"),
]
for i, (req, desc) in enumerate(nfr_data):
    set_cell_text(nfr_table.cell(i, 0), req, bold=(i == 0), font_size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(nfr_table.cell(i, 1), desc, bold=(i == 0), font_size=10)
    nfr_table.cell(i, 0).width = Inches(1.3)
    nfr_table.cell(i, 1).width = Inches(4.9)
    if i == 0:
        shade_cell(nfr_table.cell(i, 0))
        shade_cell(nfr_table.cell(i, 1))
keep_table_on_one_page(nfr_table)

add_section_heading("3.3", "System Requirements")
add_centered_text("Table 3.5: Hardware Requirements", font_size=10, bold=True, space_after=4, keep_with_next=True)
hw_table = doc.add_table(rows=6, cols=2)
hw_table.style = 'Table Grid'
hw_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hw_data = [
    ("Component", "Minimum Requirement"),
    ("Processor", "Intel Core i3 or equivalent (dual-core)"),
    ("RAM", "4 GB (recommended 8 GB for faster model training)"),
    ("Storage", "500 MB for application, dependencies, and SQLite database"),
    ("GPU", "Not required (CPU-only computation)"),
    ("Network", "Internet connection for initial setup; offline operation supported after installation"),
]
for i, (comp, req) in enumerate(hw_data):
    set_cell_text(hw_table.cell(i, 0), comp, bold=(i == 0), font_size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(hw_table.cell(i, 1), req, bold=(i == 0), font_size=10)
    hw_table.cell(i, 0).width = Inches(1.3)
    hw_table.cell(i, 1).width = Inches(4.9)
    if i == 0:
        shade_cell(hw_table.cell(i, 0))
        shade_cell(hw_table.cell(i, 1))
keep_table_on_one_page(hw_table)

add_section_heading("3.4", "SDLC Model to be Used")

add_justified_text(
    "The project follows the Agile Software Development Life Cycle (SDLC) model, which emphasizes iterative "
    "development, continuous feedback, and adaptability to changing requirements. The Agile approach is "
    "particularly suitable for this project because the system integrates multiple components (authentication, "
    "data preprocessing, model training, web application, visualization, and database) that benefit from "
    "incremental integration and testing.",
    first_line_indent=1.27
)
add_bullet("Iterative Development: The system is built in sprints, with each sprint delivering a functional increment.")
add_bullet("Continuous Integration: Authentication, ML models, web UI, database, and visualizations are integrated early and often.")
add_bullet("Flexibility: Requirements can evolve as the team gains insights from initial model training results.")
add_bullet("Rapid Prototyping: Flask's lightweight architecture enables quick prototype development and user feedback cycles.")

add_section_heading("3.5", "Software Requirements")
add_centered_text("Table 3.6: Software Requirements", font_size=10, bold=True, space_after=4, keep_with_next=True)
sw_table = doc.add_table(rows=10, cols=3)
sw_table.style = 'Table Grid'
sw_table.alignment = WD_TABLE_ALIGNMENT.CENTER
sw_data = [
    ("Category", "Software", "Version / Details"),
    ("Programming Language", "Python", "3.11"),
    ("Web Framework", "Flask", "2.x (with Jinja2 templates)"),
    ("Data Processing", "pandas, NumPy", "2.x (DataFrame manipulation, CSV I/O)"),
    ("Machine Learning", "scikit-learn", "1.x (4 classification models)"),
    ("Visualization", "Chart.js, matplotlib, seaborn", "4.x (client-side charts + server-side EDA)"),
    ("Database", "SQLite", "3.x (insurance.db, serverless)"),
    ("Frontend", "Bootstrap 5", "5.x (dark theme, responsive design)"),
    ("Security", "Werkzeug", "2.x (password hashing, PBKDF2-SHA256)"),
    ("Containerization", "Docker", "20.x+ (Python 3.11-slim base)"),
]
for i, row_data in enumerate(sw_data):
    for j, val in enumerate(row_data):
        set_cell_text(sw_table.cell(i, j), val, bold=(i == 0), font_size=10,
                      align=WD_ALIGN_PARAGRAPH.CENTER if j != 2 else WD_ALIGN_PARAGRAPH.LEFT)
    if i == 0:
        for j in range(3):
            shade_cell(sw_table.cell(i, j))
    sw_table.cell(i, 0).width = Inches(1.5)
    sw_table.cell(i, 1).width = Inches(1.5)
    sw_table.cell(i, 2).width = Inches(3.2)
keep_table_on_one_page(sw_table)

# ============================================================
# CHAPTER 4 - SYSTEM DESIGN
# ============================================================
p_ch4 = add_centered_text("CHAPTER 4", font_size=18, bold=True, space_before=24, space_after=3)
p_ch4.paragraph_format.keep_with_next = True
p_ch4.paragraph_format.page_break_before = True
p_ch4t = add_centered_text("SYSTEM DESIGN", font_size=16, bold=True, space_after=10)
p_ch4t.paragraph_format.keep_with_next = True

add_section_heading("4.1", "Design Approach")

add_justified_text(
    "The car insurance claim prediction system follows the Model-View-Controller (MVC) architectural pattern "
    "with persistent storage. The Model layer comprises the pre-trained scikit-learn classification models "
    "(loaded from pickle files at startup) and the SQLite database schema (users and predictions tables). "
    "The View layer is implemented using 9 Jinja2 templates (base.html, login.html, register.html, home.html, "
    "predict.html, history.html, visualize.html, dashboard.html, about.html) all extending a common base "
    "template with Bootstrap 5 dark theme styling and consistent navigation.",
    first_line_indent=1.27
)

add_justified_text(
    "The Controller layer consists of 10 Flask route handlers managing authentication (/, /register, /login, "
    "/logout), core features (/home, /predict, /history), analytics (/visualize, /dashboard), and information "
    "(/about). Authentication routes handle user registration with password hashing and login verification. "
    "The /predict route processes the 17-field form, applies label encoding and scaling transformations "
    "matching the training preprocessing, passes the transformed input to the Gradient Boosting model, and "
    "stores the prediction result in SQLite with a confidence score.",
    first_line_indent=1.27
)

add_justified_text(
    "The ML pipeline is executed offline during model training (train_model.py), which trains all four models, "
    "evaluates them, and serializes the best model (Gradient Boosting) along with the label encoders and "
    "scaler to pickle files. At runtime, the Flask application loads these pre-trained artifacts, enabling "
    "sub-second prediction responses. This separation of training and inference ensures fast user experience "
    "while maintaining the flexibility to retrain models with updated data.",
    first_line_indent=1.27
)

add_section_heading("4.2", "System Architecture Diagram")

add_justified_text(
    "The system architecture illustrates the interaction between the user interface, Flask application layer, "
    "machine learning pipeline, SQLite database, and visualization components. The architecture follows a "
    "request-response pattern: authenticated users submit policyholder data through the web browser, Flask "
    "processes the request, the pre-trained model generates a prediction, the result is stored in SQLite, "
    "and the prediction with confidence score is rendered back to the user.",
    first_line_indent=1.27
)

add_justified_text(
    "The system has three main data flows: (1) Authentication flow \u2014 user registration stores hashed "
    "passwords in the users table, login verifies credentials and creates a session; (2) Prediction flow "
    "\u2014 form input is preprocessed (label encoded + scaled), passed to the Gradient Boosting model, "
    "result stored in predictions table, and rendered to the user; (3) Analytics flow \u2014 dashboard and "
    "visualization pages query the models_info.json file and dataset to generate Chart.js visualizations "
    "comparing all four models.",
    first_line_indent=1.27
)

add_figure(os.path.join(FIGURES_DIR, "fig_system_architecture.png"),
           "Figure 4.1: System Architecture Diagram", width=Inches(5.5))

add_section_heading("4.3", "UML Diagrams")

add_subsection_heading("4.3.1", "Use Case Diagram")
add_justified_text(
    "The use case diagram identifies two actor types: Guest (unauthenticated) and User (authenticated). "
    "Guest actors can register and log in. Authenticated users can perform the following operations: "
    "view the home dashboard with dataset statistics, input policyholder details and receive claim "
    "predictions, view prediction history, explore EDA visualizations, analyze model performance on the "
    "dashboard, view project information on the about page, and log out. The authentication boundary "
    "separates public routes from protected routes.",
    first_line_indent=1.27
)
add_figure(os.path.join(FIGURES_DIR, "fig_use_case_diagram.png"),
           "Figure 4.2: Use Case Diagram", width=Inches(5.0))

add_subsection_heading("4.3.2", "ML Pipeline Diagram")
add_justified_text(
    "The ML pipeline diagram shows the data flow from dataset loading to model deployment. The pipeline "
    "comprises six stages: (1) Data Ingestion (load Car_Insurance_Claim.csv with pandas), (2) Preprocessing "
    "(label encode 9 categorical features, standard scale 8 numeric features), (3) Train-Test Split "
    "(80/20 stratified split preserving class distribution), (4) Model Training (train Random Forest, "
    "Gradient Boosting, SVM, and Logistic Regression), (5) Evaluation (compute accuracy, precision, recall, "
    "F1-score), and (6) Model Serialization (pickle the best model, encoders, and scaler for deployment).",
    first_line_indent=1.27
)
add_figure(os.path.join(FIGURES_DIR, "fig_ml_pipeline.png"),
           "Figure 4.3: ML Pipeline Diagram", width=Inches(5.0))

add_subsection_heading("4.3.3", "Data Preprocessing Pipeline")
add_justified_text(
    "The data preprocessing pipeline details the transformation steps applied to the raw CSV data. "
    "Step 1: Label encode 9 categorical features (GENDER, RACE, DRIVING_EXPERIENCE, EDUCATION, INCOME, "
    "VEHICLE_OWNERSHIP, VEHICLE_YEAR, MARRIED, VEHICLE_TYPE) using scikit-learn's LabelEncoder, converting "
    "text categories to integer codes. Step 2: Standard scale 8 numeric features (AGE, CREDIT_SCORE, "
    "ANNUAL_MILEAGE, SPEEDING_VIOLATIONS, DUIS, PAST_ACCIDENTS, POSTAL_CODE, CHILDREN) using StandardScaler "
    "to zero mean and unit variance. Step 3: Stratified 80/20 train-test split preserving the 74%/26% "
    "class distribution of OUTCOME.",
    first_line_indent=1.27
)
add_figure(os.path.join(FIGURES_DIR, "fig_data_preprocessing.png"),
           "Figure 4.4: Data Preprocessing Pipeline", width=Inches(5.0))

add_subsection_heading("4.3.4", "Activity Diagram")
add_justified_text(
    "The activity diagram shows the complete user workflow from application access to prediction review. "
    "A new user accesses the login page, navigates to registration, creates an account with username and "
    "password, and is redirected to login. An existing user logs in with credentials, which are verified "
    "against hashed passwords in the users table. After successful authentication, the user is directed "
    "to the home page showing dataset statistics. From there, the user can navigate to predict (fill form, "
    "submit, view result), history (review past predictions), visualize (explore EDA charts), dashboard "
    "(compare model metrics), or about (project info). The user can log out at any time.",
    first_line_indent=1.27
)
add_figure(os.path.join(FIGURES_DIR, "fig_ml_pipeline.png"),
           "Figure 4.5: Activity Diagram", width=Inches(5.0))

add_section_heading("4.4", "User Interface Design")

add_justified_text(
    "The user interface is designed using Bootstrap 5's dark theme with a consistent navigation bar across "
    "all pages. The base template (base.html) defines the common layout including the navbar with links to "
    "Home, Predict, History, Visualize, Dashboard, About, and Logout. All pages use Bootstrap's card "
    "components, form controls, and grid system for responsive layout. The dark theme reduces eye strain "
    "during extended use and provides a modern, professional appearance suitable for enterprise applications.",
    first_line_indent=1.27
)

add_justified_text(
    "The prediction page (predict.html) features a structured form with 17 input fields organized in a "
    "logical layout: personal information (age, gender, race, education, income, marital status, children), "
    "driving profile (driving experience, annual mileage, speeding violations, DUIs, past accidents), "
    "vehicle details (vehicle ownership, vehicle year, vehicle type), and financial information (credit score, "
    "postal code). The results section displays the prediction outcome (Claim / No Claim) with a confidence "
    "score, color-coded for quick visual interpretation.",
    first_line_indent=1.27
)

add_section_heading("4.5", "Database Schema")

add_justified_text(
    "The SQLite database (insurance.db) contains two tables: users for authentication and predictions for "
    "tracking prediction history. The schema is designed for simplicity and efficiency, leveraging SQLite's "
    "serverless architecture to eliminate database server configuration requirements.",
    first_line_indent=1.27
)

add_centered_text("Table 4.1: Database Schema \u2013 Users Table", font_size=10, bold=True, space_after=4, keep_with_next=True)
users_table = doc.add_table(rows=6, cols=4)
users_table.style = 'Table Grid'
users_table.alignment = WD_TABLE_ALIGNMENT.CENTER
users_data = [
    ("Column", "Type", "Constraints", "Description"),
    ("id", "INTEGER", "PRIMARY KEY AUTOINCREMENT", "Unique user identifier"),
    ("username", "TEXT", "UNIQUE NOT NULL", "Login username"),
    ("password", "TEXT", "NOT NULL", "Werkzeug hashed password"),
    ("name", "TEXT", "NOT NULL", "Display name"),
    ("role", "TEXT", "DEFAULT 'user'", "User role (user/admin)"),
]
for i, row_data in enumerate(users_data):
    for j, val in enumerate(row_data):
        set_cell_text(users_table.cell(i, j), val, bold=(i == 0), font_size=10,
                      align=WD_ALIGN_PARAGRAPH.CENTER if j in [0, 1] else WD_ALIGN_PARAGRAPH.LEFT)
    if i == 0:
        for j in range(4):
            shade_cell(users_table.cell(i, j))
    users_table.cell(i, 0).width = Inches(1.0)
    users_table.cell(i, 1).width = Inches(0.9)
    users_table.cell(i, 2).width = Inches(2.3)
    users_table.cell(i, 3).width = Inches(2.0)
keep_table_on_one_page(users_table)

add_centered_text("Table 4.2: Database Schema \u2013 Predictions Table", font_size=10, bold=True, space_before=8, space_after=4, keep_with_next=True)
pred_table = doc.add_table(rows=7, cols=4)
pred_table.style = 'Table Grid'
pred_table.alignment = WD_TABLE_ALIGNMENT.CENTER
pred_data = [
    ("Column", "Type", "Constraints", "Description"),
    ("id", "INTEGER", "PRIMARY KEY AUTOINCREMENT", "Unique prediction identifier"),
    ("user_id", "INTEGER", "NOT NULL, FK \u2192 users.id", "Reference to user who made prediction"),
    ("input_data", "TEXT", "NOT NULL", "JSON-encoded policyholder input data"),
    ("prediction", "TEXT", "NOT NULL", "Prediction outcome (Claim/No Claim)"),
    ("confidence", "REAL", "NOT NULL", "Model confidence score (0.0-1.0)"),
    ("pred_date", "TEXT", "NOT NULL", "Prediction timestamp (ISO format)"),
]
for i, row_data in enumerate(pred_data):
    for j, val in enumerate(row_data):
        set_cell_text(pred_table.cell(i, j), val, bold=(i == 0), font_size=10,
                      align=WD_ALIGN_PARAGRAPH.CENTER if j in [0, 1] else WD_ALIGN_PARAGRAPH.LEFT)
    if i == 0:
        for j in range(4):
            shade_cell(pred_table.cell(i, j))
    pred_table.cell(i, 0).width = Inches(1.0)
    pred_table.cell(i, 1).width = Inches(0.9)
    pred_table.cell(i, 2).width = Inches(2.3)
    pred_table.cell(i, 3).width = Inches(2.0)
keep_table_on_one_page(pred_table)

add_centered_text("Table 4.3: Dataset Feature Descriptions", font_size=10, bold=True, space_before=8, space_after=4, keep_with_next=True)
ds_table = doc.add_table(rows=19, cols=4)
ds_table.style = 'Table Grid'
ds_table.alignment = WD_TABLE_ALIGNMENT.CENTER
ds_data = [
    ("Feature", "Type", "Values / Range", "Description"),
    ("AGE", "Numeric", "18\u201370", "Policyholder age in years"),
    ("GENDER", "Categorical", "Male / Female", "Policyholder gender"),
    ("RACE", "Categorical", "Majority / Minority", "Demographic category"),
    ("DRIVING_EXPERIENCE", "Categorical", "0-9y, 10-19y, 20-29y, 30y+", "Years of driving experience"),
    ("EDUCATION", "Categorical", "None, High School, University", "Highest education level"),
    ("INCOME", "Categorical", "Poverty, Working, Middle, Upper", "Income bracket"),
    ("CREDIT_SCORE", "Numeric", "300\u2013850", "Credit score"),
    ("VEHICLE_OWNERSHIP", "Categorical", "Own / Not Own", "Vehicle ownership status"),
    ("VEHICLE_YEAR", "Categorical", "Before 2015 / After 2015", "Vehicle manufacture period"),
    ("MARRIED", "Categorical", "Yes / No", "Marital status"),
    ("CHILDREN", "Numeric", "0\u20135", "Number of children"),
    ("POSTAL_CODE", "Numeric", "10000\u201399999", "Residential postal code"),
    ("ANNUAL_MILEAGE", "Numeric", "1000\u201325000", "Annual miles driven"),
    ("VEHICLE_TYPE", "Categorical", "Sedan / SUV / Sports", "Vehicle category"),
    ("SPEEDING_VIOLATIONS", "Numeric", "0\u201310", "Past speeding violations count"),
    ("DUIS", "Numeric", "0\u20135", "Past DUI (driving under influence) count"),
    ("PAST_ACCIDENTS", "Numeric", "0\u201310", "Past accidents count"),
]
for i, row_data in enumerate(ds_data):
    for j, val in enumerate(row_data):
        set_cell_text(ds_table.cell(i, j), val, bold=(i == 0), font_size=8,
                      align=WD_ALIGN_PARAGRAPH.CENTER if j in [1, 2] else WD_ALIGN_PARAGRAPH.LEFT)
    if i == 0:
        for j in range(4):
            shade_cell(ds_table.cell(i, j))
    ds_table.cell(i, 0).width = Inches(1.5)
    ds_table.cell(i, 1).width = Inches(0.9)
    ds_table.cell(i, 2).width = Inches(1.5)
    ds_table.cell(i, 3).width = Inches(2.3)
keep_table_on_one_page(ds_table)

add_justified_text(
    "Note: The target variable OUTCOME is a binary column (0=No Claim ~74%, 1=Claim Filed ~26%) representing "
    "whether the policyholder filed an insurance claim. The 17 features listed above, after preprocessing "
    "(label encoding + standard scaling), serve as input to all four classification models. The dataset "
    "contains 10,000 synthetic records generated using the generate_dataset.py script.",
    font_size=11, space_before=4
)

# ============================================================
# CHAPTER 5 - IMPLEMENTATION
# ============================================================
p_ch5 = add_centered_text("CHAPTER 5", font_size=18, bold=True, space_before=24, space_after=3)
p_ch5.paragraph_format.keep_with_next = True
p_ch5.paragraph_format.page_break_before = True
p_ch5t = add_centered_text("IMPLEMENTATION", font_size=16, bold=True, space_after=10)
p_ch5t.paragraph_format.keep_with_next = True

add_section_heading("5.1", "Methodologies")

add_justified_text(
    "The project follows the Agile Software Development methodology with iterative sprints, each focused "
    "on delivering a working increment of the system. The development was organized into four sprints, "
    "each lasting approximately two weeks, with regular reviews and retrospectives to ensure alignment "
    "with project objectives.",
    first_line_indent=1.27
)

add_justified_text(
    "Sprint 1: Dataset Generation and Preprocessing \u2014 This sprint focused on creating the synthetic "
    "Car Insurance Claim dataset using generate_dataset.py (10,000 records, 17 features + OUTCOME target), "
    "implementing the label encoding pipeline for 9 categorical features, standard scaling for 8 numeric "
    "features, and the stratified 80/20 train-test split. Initial exploratory data analysis was conducted "
    "to understand feature distributions and class imbalance (74% no-claim, 26% claim).",
    first_line_indent=1.27
)

add_justified_text(
    "Sprint 2: Machine Learning Models and Evaluation \u2014 This sprint implemented the four classification "
    "models (Random Forest, Gradient Boosting, SVM, Logistic Regression) using scikit-learn. The evaluation "
    "pipeline computes accuracy, precision, recall, and F1-score for each model on the test set. Model "
    "serialization using pickle was implemented to save the best model (Gradient Boosting), label encoders, "
    "and standard scaler for runtime inference. The models_info.json file was created to store performance "
    "metrics for the dashboard.",
    first_line_indent=1.27
)

add_justified_text(
    "Sprint 3: Authentication and Database \u2014 This sprint created the SQLite database schema (users and "
    "predictions tables), implemented user registration with Werkzeug password hashing, login verification, "
    "session management, and the @login_required decorator for route protection. The prediction history "
    "storage and retrieval system was built, enabling users to review past predictions.",
    first_line_indent=1.27
)

add_justified_text(
    "Sprint 4: Web Application, Visualization, and Deployment \u2014 The final sprint created the Flask "
    "application with 10 routes, developed the 9 Jinja2 templates with Bootstrap 5 dark theme, implemented "
    "the Chart.js visualizations for the visualize and dashboard pages, and created the Dockerfile for "
    "containerized deployment on port 5002. Comprehensive testing was performed across all routes and features.",
    first_line_indent=1.27
)

add_figure(os.path.join(FIGURES_DIR, "fig_ml_pipeline.png"),
           "Figure 5.1: Agile Development Model", width=Inches(5.0))

add_section_heading("5.2", "Implementation Details")

add_justified_text(
    "The machine learning pipeline is implemented in train_model.py, which is executed once to train all "
    "four models and serialize the best performer. The script loads Car_Insurance_Claim.csv into a pandas "
    "DataFrame, applies LabelEncoder to each of the 9 categorical columns (GENDER, RACE, DRIVING_EXPERIENCE, "
    "EDUCATION, INCOME, VEHICLE_OWNERSHIP, VEHICLE_YEAR, MARRIED, VEHICLE_TYPE), and StandardScaler to the "
    "8 numeric columns (AGE, CREDIT_SCORE, ANNUAL_MILEAGE, SPEEDING_VIOLATIONS, DUIS, PAST_ACCIDENTS, "
    "POSTAL_CODE, CHILDREN). The preprocessed data is split with stratify=y to preserve class distribution.",
    first_line_indent=1.27
)

add_justified_text(
    "Each model is trained with scikit-learn's fit() method: RandomForestClassifier(n_estimators=100, "
    "random_state=42), GradientBoostingClassifier(n_estimators=100, random_state=42), SVC(kernel='rbf', "
    "probability=True, random_state=42), and LogisticRegression(max_iter=1000, random_state=42). After "
    "training, accuracy, precision, recall, and F1-score are computed using sklearn.metrics functions. "
    "The best model (Gradient Boosting with 91.95% accuracy) is serialized to claim_model.pkl, and the "
    "label encoders dictionary and scaler are saved to encoders.pkl.",
    first_line_indent=1.27
)

add_justified_text(
    "At runtime, the Flask application (app.py) loads the pre-trained model and encoders at startup. When "
    "a user submits the prediction form, the /predict route handler extracts the 17 form fields, applies "
    "the same label encoding and standard scaling transformations used during training, passes the "
    "transformed feature vector to model.predict() and model.predict_proba(), and stores the result in "
    "the predictions table with the user_id, JSON-encoded input data, prediction outcome, confidence "
    "score, and timestamp.",
    first_line_indent=1.27
)

add_section_heading("5.3", "Module Description")
add_bullet("Authentication Module: Handles user registration with Werkzeug generate_password_hash(), login verification with check_password_hash(), session management using Flask sessions, and route protection with @login_required decorator. Routes: /register, /login, /logout.")
add_bullet("Prediction Module: Processes the 17-field form input, applies label encoding and standard scaling matching training preprocessing, invokes model.predict() and model.predict_proba() on the pre-trained Gradient Boosting model, stores results in SQLite, and renders the prediction with confidence score. Route: /predict.")
add_bullet("History Module: Queries the predictions table for the current user's past predictions, displays them in a Bootstrap table with input summary, outcome, confidence, and date. Route: /history.")
add_bullet("Visualization Module: Generates Chart.js visualizations from the dataset including feature distribution charts, outcome distribution pie chart, correlation analysis, and data insights. Route: /visualize.")
add_bullet("Dashboard Module: Loads model performance metrics from models_info.json and generates Chart.js bar charts comparing accuracy, precision, recall, and F1-score across all four models. Displays confusion matrix for the best model. Route: /dashboard.")
add_bullet("Home Module: Displays dataset statistics (total records, features, class distribution), recent predictions summary, and navigation cards to other features. Route: /home.")

add_section_heading("5.4", "Sample Code")

add_subsection_heading("5.4.1", "Data Preprocessing & Model Training")
add_justified_text(
    '# Data preprocessing (train_model.py)\n'
    'df = pd.read_csv("Car_Insurance_Claim.csv")\n'
    'categorical_cols = ["GENDER", "RACE", "DRIVING_EXPERIENCE",\n'
    '    "EDUCATION", "INCOME", "VEHICLE_OWNERSHIP",\n'
    '    "VEHICLE_YEAR", "MARRIED", "VEHICLE_TYPE"]\n'
    'numeric_cols = ["AGE", "CREDIT_SCORE", "ANNUAL_MILEAGE",\n'
    '    "SPEEDING_VIOLATIONS", "DUIS", "PAST_ACCIDENTS",\n'
    '    "POSTAL_CODE", "CHILDREN"]\n'
    '\n'
    'encoders = {}\n'
    'for col in categorical_cols:\n'
    '    le = LabelEncoder()\n'
    '    df[col] = le.fit_transform(df[col])\n'
    '    encoders[col] = le\n'
    '\n'
    'scaler = StandardScaler()\n'
    'df[numeric_cols] = scaler.fit_transform(df[numeric_cols])\n'
    '\n'
    'X = df.drop("OUTCOME", axis=1)\n'
    'y = df["OUTCOME"]\n'
    'X_train, X_test, y_train, y_test = train_test_split(\n'
    '    X, y, test_size=0.2, stratify=y, random_state=42)',
    font_size=9
)

add_subsection_heading("5.4.2", "Model Training and Evaluation")
add_justified_text(
    'from sklearn.metrics import accuracy_score, precision_score\n'
    'from sklearn.metrics import recall_score, f1_score\n'
    '\n'
    'models = {\n'
    '    "Random Forest": RandomForestClassifier(\n'
    '        n_estimators=100, random_state=42),\n'
    '    "Gradient Boosting": GradientBoostingClassifier(\n'
    '        n_estimators=100, random_state=42),\n'
    '    "SVM": SVC(kernel="rbf", probability=True,\n'
    '        random_state=42),\n'
    '    "Logistic Regression": LogisticRegression(\n'
    '        max_iter=1000, random_state=42),\n'
    '}\n'
    'results = {}\n'
    'for name, model in models.items():\n'
    '    model.fit(X_train, y_train)\n'
    '    y_pred = model.predict(X_test)\n'
    '    results[name] = {\n'
    '        "accuracy": accuracy_score(y_test, y_pred),\n'
    '        "precision": precision_score(y_test, y_pred),\n'
    '        "recall": recall_score(y_test, y_pred),\n'
    '        "f1": f1_score(y_test, y_pred),\n'
    '    }',
    font_size=9
)

add_subsection_heading("5.4.3", "User Authentication Code")
add_justified_text(
    'from werkzeug.security import generate_password_hash\n'
    'from werkzeug.security import check_password_hash\n'
    '\n'
    '@app.route("/register", methods=["GET", "POST"])\n'
    'def register():\n'
    '    if request.method == "POST":\n'
    '        username = request.form["username"]\n'
    '        password = request.form["password"]\n'
    '        name = request.form["name"]\n'
    '        hashed = generate_password_hash(password)\n'
    '        db = get_db()\n'
    '        try:\n'
    '            db.execute(\n'
    '                "INSERT INTO users (username, password, name)"\n'
    '                " VALUES (?, ?, ?)",\n'
    '                (username, hashed, name))\n'
    '            db.commit()\n'
    '            flash("Registration successful!")\n'
    '            return redirect(url_for("login"))\n'
    '        except sqlite3.IntegrityError:\n'
    '            flash("Username already exists.")\n'
    '    return render_template("register.html")',
    font_size=9
)

add_subsection_heading("5.4.4", "Prediction Route Handler")
add_justified_text(
    '@app.route("/predict", methods=["GET", "POST"])\n'
    'def predict():\n'
    '    if "user_id" not in session:\n'
    '        return redirect(url_for("login"))\n'
    '    if request.method == "POST":\n'
    '        input_data = {}\n'
    '        for field in feature_names:\n'
    '            input_data[field] = request.form[field]\n'
    '        # Apply label encoding and scaling\n'
    '        features = preprocess_input(input_data)\n'
    '        prediction = model.predict([features])[0]\n'
    '        proba = model.predict_proba([features])[0]\n'
    '        confidence = max(proba)\n'
    '        result = "Claim" if prediction == 1 else "No Claim"\n'
    '        # Store in database\n'
    '        db = get_db()\n'
    '        db.execute(\n'
    '            "INSERT INTO predictions "\n'
    '            "(user_id, input_data, prediction, confidence,"\n'
    '            " pred_date) VALUES (?, ?, ?, ?, ?)",\n'
    '            (session["user_id"], json.dumps(input_data),\n'
    '             result, confidence, datetime.now().isoformat()))\n'
    '        db.commit()\n'
    '        return render_template("predict.html",\n'
    '            result=result, confidence=confidence)\n'
    '    return render_template("predict.html")',
    font_size=9
)

add_subsection_heading("5.4.5", "Database Initialization")
add_justified_text(
    'import sqlite3\n'
    '\n'
    'def init_db():\n'
    '    db = sqlite3.connect("insurance.db")\n'
    '    db.execute("""\n'
    '        CREATE TABLE IF NOT EXISTS users (\n'
    '            id INTEGER PRIMARY KEY AUTOINCREMENT,\n'
    '            username TEXT UNIQUE NOT NULL,\n'
    '            password TEXT NOT NULL,\n'
    '            name TEXT NOT NULL,\n'
    '            role TEXT DEFAULT \'user\')\n'
    '    """)\n'
    '    db.execute("""\n'
    '        CREATE TABLE IF NOT EXISTS predictions (\n'
    '            id INTEGER PRIMARY KEY AUTOINCREMENT,\n'
    '            user_id INTEGER NOT NULL,\n'
    '            input_data TEXT NOT NULL,\n'
    '            prediction TEXT NOT NULL,\n'
    '            confidence REAL NOT NULL,\n'
    '            pred_date TEXT NOT NULL,\n'
    '            FOREIGN KEY (user_id) REFERENCES users(id))\n'
    '    """)\n'
    '    db.commit()\n'
    '    db.close()',
    font_size=9
)

# ============================================================
# CHAPTER 6 - TESTING
# ============================================================
p_ch6 = add_centered_text("CHAPTER 6", font_size=18, bold=True, space_before=24, space_after=3)
p_ch6.paragraph_format.keep_with_next = True
p_ch6.paragraph_format.page_break_before = True
p_ch6t = add_centered_text("TESTING", font_size=16, bold=True, space_after=10)
p_ch6t.paragraph_format.keep_with_next = True

add_section_heading("6.1", "Types of Testing")

add_subsection_heading("6.1.1", "Unit Testing")
add_justified_text(
    "Unit testing verifies the correct behavior of individual components in isolation. For the car insurance "
    "claim prediction system, unit tests cover the data preprocessing functions (verifying label encoding "
    "produces valid integer codes, standard scaling produces zero-mean unit-variance outputs), the model "
    "prediction pipeline (verifying that the model produces binary predictions 0 or 1), the authentication "
    "functions (verifying password hashing and verification), and the database operations (verifying insert "
    "and query operations on users and predictions tables).",
    first_line_indent=1.27
)

add_subsection_heading("6.1.2", "Integration Testing")
add_justified_text(
    "Integration testing validates the interaction between connected components. Key integration tests include "
    "the end-to-end prediction pipeline (form submission \u2192 preprocessing \u2192 model inference \u2192 "
    "database storage \u2192 result rendering), the authentication flow (registration \u2192 login \u2192 "
    "session creation \u2192 protected route access \u2192 logout), and the history retrieval pipeline "
    "(prediction storage \u2192 history query \u2192 table rendering). These tests ensure that data flows "
    "correctly between Flask routes, the ML pipeline, SQLite database, and the template layer.",
    first_line_indent=1.27
)

add_subsection_heading("6.1.3", "Functional Testing")
add_justified_text(
    "Functional testing evaluates the system against its specified requirements. This includes testing user "
    "registration with valid and invalid inputs, login with correct and incorrect credentials, prediction "
    "with all 17 fields filled correctly, history page displaying past predictions, visualization page "
    "rendering Chart.js charts, dashboard page showing model comparison metrics, and error handling for "
    "duplicate usernames, empty forms, and unauthorized access. Each functional test verifies that the "
    "user-facing behavior matches the specification.",
    first_line_indent=1.27
)

add_section_heading("6.2", "Test Cases")

add_centered_text("Table 6.1: Test Cases \u2013 Authentication", font_size=10, bold=True, space_after=4, keep_with_next=True)
tc1 = doc.add_table(rows=6, cols=4)
tc1.style = 'Table Grid'
tc1.alignment = WD_TABLE_ALIGNMENT.CENTER
tc1_data = [
    ("Test ID", "Scenario", "Expected Result", "Status"),
    ("TC-A01", "Register with valid username, password, name", "Account created; redirect to login page", "Pass"),
    ("TC-A02", "Register with duplicate username", "Error: Username already exists", "Pass"),
    ("TC-A03", "Login with correct credentials", "Session created; redirect to home page", "Pass"),
    ("TC-A04", "Login with incorrect password", "Error: Invalid credentials", "Pass"),
    ("TC-A05", "Access protected route without login", "Redirect to login page", "Pass"),
]
for i, row_data in enumerate(tc1_data):
    for j, val in enumerate(row_data):
        set_cell_text(tc1.cell(i, j), val, bold=(i == 0), font_size=9,
                      align=WD_ALIGN_PARAGRAPH.CENTER if j in [0, 3] else WD_ALIGN_PARAGRAPH.LEFT)
    if i == 0:
        for j in range(4):
            shade_cell(tc1.cell(i, j))
    tc1.cell(i, 0).width = Inches(0.7)
    tc1.cell(i, 1).width = Inches(2.0)
    tc1.cell(i, 2).width = Inches(2.5)
    tc1.cell(i, 3).width = Inches(0.6)
keep_table_on_one_page(tc1)

add_centered_text("Table 6.2: Test Cases \u2013 Prediction", font_size=10, bold=True, space_before=8, space_after=4, keep_with_next=True)
tc2 = doc.add_table(rows=6, cols=4)
tc2.style = 'Table Grid'
tc2.alignment = WD_TABLE_ALIGNMENT.CENTER
tc2_data = [
    ("Test ID", "Scenario", "Expected Result", "Status"),
    ("TC-P01", "Submit prediction with all 17 fields valid", "Prediction displayed with confidence score; stored in DB", "Pass"),
    ("TC-P02", "Submit prediction for low-risk profile", "No Claim prediction with high confidence (>80%)", "Pass"),
    ("TC-P03", "Submit prediction for high-risk profile", "Claim prediction with appropriate confidence", "Pass"),
    ("TC-P04", "View prediction history after multiple predictions", "All past predictions listed with correct details", "Pass"),
    ("TC-P05", "Verify prediction stored in SQLite database", "predictions table contains correct user_id, input, result", "Pass"),
]
for i, row_data in enumerate(tc2_data):
    for j, val in enumerate(row_data):
        set_cell_text(tc2.cell(i, j), val, bold=(i == 0), font_size=9,
                      align=WD_ALIGN_PARAGRAPH.CENTER if j in [0, 3] else WD_ALIGN_PARAGRAPH.LEFT)
    if i == 0:
        for j in range(4):
            shade_cell(tc2.cell(i, j))
    tc2.cell(i, 0).width = Inches(0.7)
    tc2.cell(i, 1).width = Inches(2.0)
    tc2.cell(i, 2).width = Inches(2.5)
    tc2.cell(i, 3).width = Inches(0.6)
keep_table_on_one_page(tc2)

add_centered_text("Table 6.3: Test Cases \u2013 Visualization & Dashboard", font_size=10, bold=True, space_before=8, space_after=4, keep_with_next=True)
tc3 = doc.add_table(rows=5, cols=4)
tc3.style = 'Table Grid'
tc3.alignment = WD_TABLE_ALIGNMENT.CENTER
tc3_data = [
    ("Test ID", "Scenario", "Expected Result", "Status"),
    ("TC-V01", "Render EDA visualization page", "Chart.js charts displayed for feature distributions", "Pass"),
    ("TC-V02", "Render model dashboard", "Bar charts showing accuracy, precision, recall, F1 for all 4 models", "Pass"),
    ("TC-V03", "Render confusion matrix on dashboard", "Confusion matrix displayed for Gradient Boosting model", "Pass"),
    ("TC-V04", "Dashboard loads model metrics from JSON", "All 4 model metrics loaded and displayed correctly", "Pass"),
]
for i, row_data in enumerate(tc3_data):
    for j, val in enumerate(row_data):
        set_cell_text(tc3.cell(i, j), val, bold=(i == 0), font_size=9,
                      align=WD_ALIGN_PARAGRAPH.CENTER if j in [0, 3] else WD_ALIGN_PARAGRAPH.LEFT)
    if i == 0:
        for j in range(4):
            shade_cell(tc3.cell(i, j))
    tc3.cell(i, 0).width = Inches(0.7)
    tc3.cell(i, 1).width = Inches(2.0)
    tc3.cell(i, 2).width = Inches(2.5)
    tc3.cell(i, 3).width = Inches(0.6)
keep_table_on_one_page(tc3)

add_centered_text("Table 6.4: Test Cases \u2013 Data & Security", font_size=10, bold=True, space_before=8, space_after=4, keep_with_next=True)
tc4 = doc.add_table(rows=5, cols=4)
tc4.style = 'Table Grid'
tc4.alignment = WD_TABLE_ALIGNMENT.CENTER
tc4_data = [
    ("Test ID", "Scenario", "Expected Result", "Status"),
    ("TC-S01", "Verify passwords stored as hashes", "No plain-text passwords in users table", "Pass"),
    ("TC-S02", "Verify session cleared on logout", "Protected routes inaccessible after logout", "Pass"),
    ("TC-S03", "Label encoding matches training encoding", "Same integer codes applied during inference as training", "Pass"),
    ("TC-S04", "Standard scaling matches training scaler", "Same mean and std applied during inference as training", "Pass"),
]
for i, row_data in enumerate(tc4_data):
    for j, val in enumerate(row_data):
        set_cell_text(tc4.cell(i, j), val, bold=(i == 0), font_size=9,
                      align=WD_ALIGN_PARAGRAPH.CENTER if j in [0, 3] else WD_ALIGN_PARAGRAPH.LEFT)
    if i == 0:
        for j in range(4):
            shade_cell(tc4.cell(i, j))
    tc4.cell(i, 0).width = Inches(0.7)
    tc4.cell(i, 1).width = Inches(2.0)
    tc4.cell(i, 2).width = Inches(2.5)
    tc4.cell(i, 3).width = Inches(0.6)
keep_table_on_one_page(tc4)

# ============================================================
# CHAPTER 7 - RESULTS AND DISCUSSION
# ============================================================
p_ch7 = add_centered_text("CHAPTER 7", font_size=18, bold=True, space_before=24, space_after=3)
p_ch7.paragraph_format.keep_with_next = True
p_ch7.paragraph_format.page_break_before = True
p_ch7t = add_centered_text("RESULTS AND DISCUSSION", font_size=16, bold=True, space_after=10)
p_ch7t.paragraph_format.keep_with_next = True

add_justified_text(
    "This chapter presents the results of the car insurance claim prediction system through application "
    "screenshots demonstrating the user interface and functionality, followed by machine learning model "
    "performance figures. The system was tested with the synthetic Car Insurance Claim dataset (10,000 rows, "
    "17 features + OUTCOME) to validate the prediction accuracy, model comparison, and visualization quality.",
    first_line_indent=1.27
)

# --- Application Screenshots ---
add_section_heading("7.1", "Login Page")
add_justified_text(
    "The login page presents a Bootstrap 5 dark-themed interface with a centered card containing username "
    "and password input fields. The page includes a link to the registration page for new users. Error "
    "messages are displayed using Flask flash messages when invalid credentials are entered. The dark theme "
    "provides a professional, modern appearance consistent across all application pages.",
    first_line_indent=1.27
)
add_figure(os.path.join(SCREENSHOTS_DIR, "login.png"),
           "Figure 7.1: Login Page", width=Inches(5.5))

add_section_heading("7.2", "Registration Page")
add_justified_text(
    "The registration page allows new users to create an account by providing a username, password, and "
    "display name. The form validates input and checks for duplicate usernames before creating the account "
    "with a Werkzeug-hashed password stored in the SQLite users table. Upon successful registration, the "
    "user is redirected to the login page with a success flash message.",
    first_line_indent=1.27
)
add_figure(os.path.join(SCREENSHOTS_DIR, "register.png"),
           "Figure 7.2: Registration Page", width=Inches(5.5))

add_section_heading("7.3", "Invalid Login Attempt")
add_justified_text(
    "When a user enters incorrect credentials, the system displays an error message indicating invalid "
    "username or password. The password verification uses Werkzeug's check_password_hash() function, "
    "which compares the entered password against the stored PBKDF2-SHA256 hash. The error message is "
    "generic (not revealing whether the username or password is wrong) to prevent enumeration attacks.",
    first_line_indent=1.27
)
add_figure(os.path.join(SCREENSHOTS_DIR, "invalid_login.png"),
           "Figure 7.3: Invalid Login Attempt", width=Inches(5.5))

add_section_heading("7.4", "Duplicate Registration Attempt")
add_justified_text(
    "When a user attempts to register with a username that already exists in the database, the system "
    "catches the SQLite IntegrityError (from the UNIQUE constraint on the username column) and displays "
    "an appropriate error message. This prevents duplicate accounts and maintains data integrity in the "
    "users table.",
    first_line_indent=1.27
)
add_figure(os.path.join(SCREENSHOTS_DIR, "duplicate_register.png"),
           "Figure 7.4: Duplicate Registration Attempt", width=Inches(5.5))

add_section_heading("7.5", "Home Page Dashboard")
add_justified_text(
    "After successful login, the user is directed to the home page displaying dataset statistics (total "
    "records, number of features, class distribution), navigation cards linking to Predict, History, "
    "Visualize, Dashboard, and About pages, and a welcome message with the user's display name. The "
    "home page provides a central hub for navigating all application features.",
    first_line_indent=1.27
)
add_figure(os.path.join(SCREENSHOTS_DIR, "home.png"),
           "Figure 7.5: Home Page Dashboard", width=Inches(5.5))

add_section_heading("7.6", "Prediction Form")
add_justified_text(
    "The prediction page features a structured form with 17 input fields organized logically: personal "
    "information (age, gender, race, education, income, marital status, children), driving profile "
    "(driving experience, annual mileage, speeding violations, DUIs, past accidents), vehicle details "
    "(vehicle ownership, vehicle year, vehicle type), and financial information (credit score, postal code). "
    "Dropdown menus are used for categorical fields and number inputs for numeric fields.",
    first_line_indent=1.27
)
add_figure(os.path.join(SCREENSHOTS_DIR, "predict_form.png"),
           "Figure 7.6: Prediction Form (Empty)", width=Inches(5.5))

add_section_heading("7.7", "Prediction Form (Filled)")
add_justified_text(
    "This screenshot shows the prediction form filled with sample policyholder data. The user has entered "
    "values for all 17 fields representing a specific insurance policy scenario. Upon clicking the Predict "
    "button, the form data is submitted to the /predict POST route for preprocessing and model inference.",
    first_line_indent=1.27
)
add_figure(os.path.join(SCREENSHOTS_DIR, "predict_filled.png"),
           "Figure 7.7: Prediction Form (Filled)", width=Inches(5.5))

add_section_heading("7.8", "Prediction Result")
add_justified_text(
    "The prediction result is displayed prominently showing the outcome (Claim or No Claim) with a "
    "confidence score. The result is color-coded \u2014 green for No Claim and red for Claim \u2014 for "
    "quick visual interpretation. The confidence score represents the probability from the Gradient Boosting "
    "model's predict_proba() output, indicating the model's certainty in its prediction. The prediction "
    "is automatically stored in the SQLite database for future reference.",
    first_line_indent=1.27
)
add_figure(os.path.join(SCREENSHOTS_DIR, "predict_result.png"),
           "Figure 7.8: Prediction Result", width=Inches(5.5))

add_section_heading("7.9", "Prediction History")
add_justified_text(
    "The history page displays a Bootstrap table listing all past predictions made by the current user. "
    "Each row shows a summary of the input data, the prediction outcome, confidence score, and the date "
    "and time of the prediction. The table is sorted by date in descending order (newest first). This "
    "feature enables users to track prediction patterns and review past risk assessments.",
    first_line_indent=1.27
)
add_figure(os.path.join(SCREENSHOTS_DIR, "history.png"),
           "Figure 7.9: Prediction History", width=Inches(5.5))

add_section_heading("7.10", "EDA Visualization Gallery")
add_justified_text(
    "The visualization page presents a gallery of Chart.js charts providing exploratory data analysis of "
    "the Car Insurance Claim dataset. Charts include feature distribution histograms, outcome distribution "
    "pie chart (74% No Claim, 26% Claim), and correlation analyses between key features. The interactive "
    "charts support hover tooltips showing exact values and responsive resizing for different screen sizes.",
    first_line_indent=1.27
)
add_figure(os.path.join(SCREENSHOTS_DIR, "visualize.png"),
           "Figure 7.10: EDA Visualization Gallery", width=Inches(5.5))

add_section_heading("7.11", "Model Dashboard")
add_justified_text(
    "The dashboard page displays comprehensive model analytics comparing all four classification models. "
    "The main chart shows grouped bar charts with accuracy, precision, recall, and F1-score for Random "
    "Forest, Gradient Boosting, SVM, and Logistic Regression. Metric cards highlight the best model "
    "(Gradient Boosting with 91.95% accuracy) and provide a quick comparison summary.",
    first_line_indent=1.27
)
add_figure(os.path.join(SCREENSHOTS_DIR, "dashboard.png"),
           "Figure 7.11: Model Dashboard", width=Inches(5.5))

add_section_heading("7.12", "Dashboard Charts (Detailed)")
add_justified_text(
    "The detailed dashboard view shows additional Chart.js visualizations including a confusion matrix "
    "for the Gradient Boosting model, a radar chart comparing model performance dimensions, and precision-"
    "recall tradeoff analysis. These charts provide deeper insights into model behavior beyond simple "
    "accuracy metrics, helping users understand where each model excels and where it struggles.",
    first_line_indent=1.27
)
add_figure(os.path.join(SCREENSHOTS_DIR, "dashboard_charts.png"),
           "Figure 7.12: Dashboard Charts (Detailed)", width=Inches(5.5))

add_section_heading("7.13", "About Page")
add_justified_text(
    "The about page provides project information including the project title, description, technology stack "
    "(Python, Flask, scikit-learn, SQLite, Bootstrap 5, Chart.js, Docker), dataset details, and team "
    "information. The page uses Bootstrap cards to organize information sections and includes links to "
    "relevant documentation and resources.",
    first_line_indent=1.27
)
add_figure(os.path.join(SCREENSHOTS_DIR, "about.png"),
           "Figure 7.13: About Page", width=Inches(5.5))

# --- Model Performance Figures ---
add_section_heading("7.14", "Model Accuracy Comparison")
add_justified_text(
    "The model accuracy comparison figure provides a detailed view of each model's classification accuracy. "
    "Gradient Boosting (91.95%) leads, followed by SVM (91.10%), Logistic Regression (90.25%), and Random "
    "Forest (90.10%). The relatively close accuracy scores indicate that all four models effectively learn "
    "the claim prediction patterns, but Gradient Boosting's sequential optimization provides a consistent "
    "edge across all metrics.",
    first_line_indent=1.27
)
add_figure(os.path.join(FIGURES_DIR, "fig_model_comparison.png"),
           "Figure 7.14: Model Accuracy Comparison", width=Inches(5.5))

add_section_heading("7.15", "Confusion Matrix")
add_justified_text(
    "The confusion matrix for the Gradient Boosting model shows the breakdown of predictions: true negatives "
    "(correctly predicted No Claim), false positives (incorrectly predicted Claim), false negatives "
    "(missed actual Claims), and true positives (correctly predicted Claims). The matrix reveals that the "
    "model is stronger at identifying no-claim cases (high specificity) than detecting actual claims "
    "(moderate sensitivity), reflecting the class imbalance in the dataset.",
    first_line_indent=1.27
)
add_figure(os.path.join(FIGURES_DIR, "fig_confusion_matrix.png"),
           "Figure 7.15: Confusion Matrix (Gradient Boosting)", width=Inches(5.5))

add_section_heading("7.16", "Feature Importance")
add_justified_text(
    "The feature importance chart from the Gradient Boosting model reveals the most influential predictors "
    "of insurance claim likelihood. Driving experience, past accidents, and speeding violations rank as the "
    "top three most important features, consistent with domain knowledge that driving behavior strongly "
    "predicts claim risk. Credit score, annual mileage, and DUIs also show significant importance. These "
    "feature importance rankings provide actionable insights for insurance underwriters, confirming that "
    "driving history and behavior are stronger predictors than demographic factors.",
    first_line_indent=1.27
)
add_figure(os.path.join(FIGURES_DIR, "fig_feature_importance.png"),
           "Figure 7.16: Feature Importance", width=Inches(5.5))

add_section_heading("7.17", "System Architecture")
add_justified_text(
    "The system architecture figure provides a high-level view of the component interactions. The user "
    "layer (web browser) communicates with the Flask application layer through HTTP requests. The Flask "
    "layer coordinates four major subsystems: the authentication module (Werkzeug + SQLite), the ML "
    "prediction pipeline (pre-trained Gradient Boosting model), the Chart.js visualization engine, and "
    "the SQLite database for persistent storage. Session management connects the authentication and "
    "prediction subsystems.",
    first_line_indent=1.27
)
add_figure(os.path.join(FIGURES_DIR, "fig_system_architecture.png"),
           "Figure 7.17: System Architecture", width=Inches(5.5))

add_section_heading("7.18", "Data Preprocessing Pipeline")
add_justified_text(
    "The data preprocessing pipeline figure illustrates the transformation steps applied to the raw dataset. "
    "Step 1: Label encode 9 categorical features using LabelEncoder, converting text categories to integer "
    "codes. Step 2: Standard scale 8 numeric features using StandardScaler, normalizing to zero mean and "
    "unit variance. Step 3: Stratified 80/20 train-test split preserving the 74%/26% class distribution. "
    "The same encoding and scaling transformations (stored in encoders.pkl) are applied during runtime "
    "inference to ensure consistency between training and prediction.",
    first_line_indent=1.27
)
add_figure(os.path.join(FIGURES_DIR, "fig_data_preprocessing.png"),
           "Figure 7.18: Data Preprocessing Pipeline", width=Inches(5.5))

# Model Performance Table
add_centered_text("Table 7.1: Model Performance Comparison", font_size=10, bold=True, space_before=8, space_after=4, keep_with_next=True)
perf_table = doc.add_table(rows=5, cols=5)
perf_table.style = 'Table Grid'
perf_table.alignment = WD_TABLE_ALIGNMENT.CENTER
perf_data = [
    ("Model", "Accuracy", "Precision", "Recall", "F1-Score"),
    ("Random Forest (n=100)", "90.10%", "86.10%", "73.85%", "79.50%"),
    ("Gradient Boosting (BEST)", "91.95%", "89.45%", "78.27%", "83.49%"),
    ("SVM (RBF kernel)", "91.10%", "85.62%", "79.04%", "82.20%"),
    ("Logistic Regression", "90.25%", "83.92%", "77.31%", "80.48%"),
]
for i, row_data in enumerate(perf_data):
    for j, val in enumerate(row_data):
        set_cell_text(perf_table.cell(i, j), val, bold=(i == 0), font_size=10,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
    if i == 0:
        for j in range(5):
            shade_cell(perf_table.cell(i, j))
    perf_table.cell(i, 0).width = Inches(2.0)
    perf_table.cell(i, 1).width = Inches(0.9)
    perf_table.cell(i, 2).width = Inches(0.9)
    perf_table.cell(i, 3).width = Inches(0.9)
    perf_table.cell(i, 4).width = Inches(0.9)
keep_table_on_one_page(perf_table)

# ============================================================
# CHAPTER 8 - CONCLUSION AND FUTURE SCOPE
# ============================================================
p_ch8 = add_centered_text("CHAPTER 8", font_size=18, bold=True, space_before=24, space_after=3)
p_ch8.paragraph_format.keep_with_next = True
p_ch8.paragraph_format.page_break_before = True
p_ch8t = add_centered_text("CONCLUSION AND FUTURE SCOPE", font_size=16, bold=True, space_after=10)
p_ch8t.paragraph_format.keep_with_next = True

add_section_heading("8.1", "Conclusion")

add_justified_text(
    "This project successfully demonstrates the application of machine learning classification techniques "
    "for predicting car insurance claim outcomes from policyholder demographics and driving history. The "
    "web-based platform implements four classification models \u2014 Random Forest, Gradient Boosting, SVM, "
    "and Logistic Regression \u2014 providing comprehensive performance comparison and secure prediction "
    "capabilities with user authentication and history tracking.",
    first_line_indent=1.27
)

add_justified_text(
    "The comprehensive comparison of four classification algorithms demonstrates clear performance "
    "differentiation. Gradient Boosting achieved the best overall performance with 91.95% accuracy, "
    "89.45% precision, 78.27% recall, and 83.49% F1-score, followed by SVM (91.10% accuracy), Logistic "
    "Regression (90.25% accuracy), and Random Forest (90.10% accuracy). The ensemble methods' ability to "
    "capture non-linear interactions between policyholder features and claim outcomes is the key factor "
    "in their superior performance.",
    first_line_indent=1.27
)

add_justified_text(
    "The Flask-based web application provides a practical, accessible interface that enables insurance "
    "professionals and analysts to perform advanced claim risk assessment without programming knowledge. "
    "The following key achievements were realized:",
    first_line_indent=1.27
)

add_bullet("Implemented four classification models achieving accuracy scores ranging from 90.10% (Random Forest) to 91.95% (Gradient Boosting) on the Car Insurance Claim dataset.")
add_bullet("Developed automated data preprocessing pipeline handling label encoding of 9 categorical features and standard scaling of 8 numeric features.")
add_bullet("Built secure user authentication system with Werkzeug password hashing and SQLite database storage.")
add_bullet("Created prediction history tracking enabling users to review and analyze past risk assessments.")
add_bullet("Implemented interactive Chart.js visualizations for exploratory data analysis and model comparison dashboards.")
add_bullet("Designed a comprehensive 9-page web application with Bootstrap 5 dark theme for professional appearance.")
add_bullet("Containerized the application with Docker for reproducible deployment on port 5002.")

add_justified_text(
    "The feature importance analysis revealed that driving experience, past accidents, and speeding violations "
    "are the strongest predictors of insurance claim likelihood, providing actionable insights for insurance "
    "underwriters. The project demonstrates that machine learning can effectively augment traditional "
    "actuarial approaches for claim risk assessment, making advanced analytics accessible to a broader "
    "audience through intuitive web interfaces with secure authentication.",
    first_line_indent=1.27
)

add_section_heading("8.2", "Future Scope")

add_justified_text(
    "The car insurance claim prediction system can be extended in several directions to enhance its "
    "analytical capabilities and practical utility:",
    first_line_indent=1.27
)
add_bullet("Deep Learning Models: Integrate neural network classifiers (multi-layer perceptrons, attention-based architectures) that may capture more complex feature interactions than tree-based methods.")
add_bullet("Claim Amount Prediction: Extend the system from binary classification (claim/no-claim) to regression predicting the actual claim amount, enabling more precise premium pricing.")
add_bullet("Real-time Data Integration: Connect to insurance company databases and telematics devices for real-time risk assessment based on actual driving behavior data.")
add_bullet("Advanced Feature Engineering: Implement interaction terms, polynomial features, and domain-specific derived variables (e.g., risk score composites) to improve prediction accuracy.")
add_bullet("Cloud Deployment: Deploy on cloud platforms (AWS, Azure, GCP) with auto-scaling and managed databases for enterprise-level concurrent user support.")
add_bullet("REST API: Develop RESTful API endpoints for integration with existing insurance management systems and third-party applications.")
add_bullet("Explainability: Integrate SHAP (SHapley Additive exPlanations) for model-agnostic feature importance, providing individual prediction explanations for regulatory compliance.")
add_bullet("Fraud Detection: Extend the system to include a fraud detection module that identifies suspicious claim patterns using anomaly detection algorithms.")
add_bullet("Multi-class Classification: Support multiple claim categories (minor, moderate, major, total loss) instead of binary claim/no-claim classification.")
add_bullet("Admin Dashboard: Add administrator role with capabilities to manage users, view all predictions, export reports, and retrain models with updated data.")

# ============================================================
# CHAPTER 9 - SUSTAINABLE DEVELOPMENT GOALS
# ============================================================
p_ch9 = add_centered_text("CHAPTER 9", font_size=18, bold=True, space_before=24, space_after=3)
p_ch9.paragraph_format.keep_with_next = True
p_ch9.paragraph_format.page_break_before = True
p_ch9t = add_centered_text("SUSTAINABLE DEVELOPMENT GOALS", font_size=16, bold=True, space_after=10)
p_ch9t.paragraph_format.keep_with_next = True

add_section_heading("9.1", "SDG 8: Decent Work and Economic Growth")

add_justified_text(
    "The car insurance claim prediction system directly contributes to SDG 8 (Decent Work and Economic "
    "Growth) by providing data-driven tools that optimize insurance industry operations. Accurate claim "
    "prediction enables insurance companies to set fair premium rates, maintain adequate reserves, and "
    "reduce operational costs associated with claims processing. By automating risk assessment through "
    "machine learning, the system reduces the manual labor required for underwriting decisions while "
    "improving the accuracy and consistency of these assessments.",
    first_line_indent=1.27
)
add_justified_text(
    "The insurance industry is a significant contributor to economic growth, providing financial protection "
    "that enables individuals and businesses to take calculated risks. More accurate claim prediction "
    "contributes to the financial sustainability of insurance companies, protecting policyholder funds and "
    "maintaining industry stability. The system's feature importance analysis reveals which risk factors "
    "most strongly predict claims, enabling evidence-based policy design that promotes responsible driving "
    "behavior and contributes to economic productivity through reduced accident rates.",
    first_line_indent=1.27
)

add_section_heading("9.2", "SDG 9: Industry, Innovation and Infrastructure")

add_justified_text(
    "The project advances insurance technology infrastructure by demonstrating how machine learning can be "
    "packaged into accessible web applications that democratize data analysis capabilities. The open-source "
    "technology stack (Python, Flask, scikit-learn, SQLite, Bootstrap 5, Docker) demonstrates that "
    "sophisticated analytical tools can be built without proprietary software or expensive licensing, "
    "making them accessible to insurance companies of all sizes, including small and medium enterprises "
    "in developing economies.",
    first_line_indent=1.27
)
add_justified_text(
    "The Docker containerization approach ensures that the system can be deployed consistently across "
    "different computing environments, from local workstations to cloud servers, supporting the "
    "development of resilient digital infrastructure for the insurance industry. The system's modular "
    "architecture (separating training, inference, authentication, and visualization) establishes a "
    "replicable pattern for building AI-powered insurance technology applications.",
    first_line_indent=1.27
)

add_section_heading("9.3", "SDG 16: Peace, Justice and Strong Institutions")

add_justified_text(
    "The car insurance claim prediction system contributes to SDG 16 (Peace, Justice and Strong "
    "Institutions) by promoting fair and transparent insurance claim assessment. Traditional manual "
    "underwriting can be influenced by unconscious biases, leading to unfair premium pricing or claim "
    "decisions. Machine learning models, when trained on comprehensive data, provide consistent and "
    "objective risk assessments based on measurable factors (driving history, vehicle characteristics, "
    "experience) rather than subjective judgments.",
    first_line_indent=1.27
)
add_justified_text(
    "The system's feature importance transparency reveals which factors drive claim predictions, enabling "
    "regulators and consumers to understand and audit the decision-making process. This transparency "
    "supports accountability in insurance pricing and contributes to building trust in financial "
    "institutions. The secure authentication and data storage practices implemented in the system also "
    "demonstrate responsible data governance, protecting policyholder information while enabling "
    "analytical insights.",
    first_line_indent=1.27
)

add_section_heading("9.4", "Broader Impact")

add_bullet("Environmental Impact: Accurate claim prediction can indirectly reduce environmental impact by incentivizing safer driving (reducing accidents and associated vehicle damage/waste). The lightweight computational requirements (CPU-only, pre-trained model inference) minimize energy consumption.")
add_bullet("Social Impact: By making ML-based risk assessment accessible through a web interface, the system enables insurance professionals in developing regions to perform advanced analyses without specialized training. Fair claim prediction promotes equitable access to insurance services.")
add_bullet("Economic Impact: The free, open-source technology stack eliminates software licensing costs for insurance companies. Accurate claim prediction reduces unexpected losses, stabilizes premium pricing, and contributes to overall insurance market efficiency.")
add_bullet("Technological Impact: The project demonstrates the feasibility of deploying ML classification pipelines as secure web applications with authentication and persistent storage, establishing a replicable pattern for other financial analytics applications.")

add_section_heading("9.5", "Future Contribution to SDGs")

add_bullet("SDG 3 (Good Health and Well-Being): Extending the system to analyze the relationship between claim patterns and injury severity could contribute to road safety insights and public health policy.")
add_bullet("SDG 10 (Reduced Inequalities): Implementing bias detection and fairness metrics to ensure that ML models do not discriminate based on protected characteristics would support equitable insurance access.")
add_bullet("SDG 17 (Partnerships): Collaborating with insurance regulators, industry associations, and academic institutions for model validation and standardization would strengthen the system's real-world impact and create cross-sector partnerships.")

# ============================================================
# REFERENCES
# ============================================================
p_ref = add_centered_text("REFERENCES", font_size=18, bold=True, space_before=24, space_after=12)
p_ref.paragraph_format.page_break_before = True

references = [
    '[1] Breiman, L. (2001). "Random Forests." Machine Learning, 45(1), 5-32.',
    '[2] Friedman, J. H. (2001). "Greedy Function Approximation: A Gradient Boosting Machine." Annals of Statistics, 29(5), 1189-1232.',
    '[3] Cortes, C., & Vapnik, V. (1995). "Support-Vector Networks." Machine Learning, 20(3), 273-297.',
    '[4] Hosmer, D. W., Lemeshow, S., & Sturdivant, R. X. (2013). "Applied Logistic Regression." Wiley, 3rd Edition.',
    '[5] Noll, A., Salzmann, R., & W\u00fcthrich, M. V. (2018). "Case Study: French Motor Third-Party Liability Claims." SSRN Electronic Journal.',
    '[6] Chen, T., & Guestrin, C. (2016). "XGBoost: A Scalable Tree Boosting System." KDD 2016.',
    '[7] Kuhn, M., & Johnson, K. (2013). "Applied Predictive Modeling." Springer.',
    '[8] He, H., & Garcia, E. A. (2009). "Learning from Imbalanced Data." IEEE Transactions on Knowledge and Data Engineering, 21(9), 1263-1284.',
    '[9] Patki, N., Wedge, R., & Veeramachaneni, K. (2016). "The Synthetic Data Vault." IEEE DSAA 2016.',
    '[10] Powers, D. M. W. (2011). "Evaluation: From Precision, Recall and F-Measure to ROC, Informedness, Markedness and Correlation." Journal of Machine Learning Technologies, 2(1), 37-63.',
    '[11] OWASP Foundation (2021). "OWASP Application Security Verification Standard 4.0." owasp.org.',
    '[12] Grinberg, M. (2018). "Flask Web Development: Developing Web Applications with Python." O\'Reilly Media.',
    '[13] Chart.js Contributors (2023). "Chart.js \u2014 Simple yet Flexible JavaScript Charting." chartjs.org.',
    '[14] Owens, M. (2006). "The Definitive Guide to SQLite." Apress.',
    '[15] Bootstrap Contributors (2021). "Bootstrap 5 Documentation." getbootstrap.com.',
    '[16] Merkel, D. (2014). "Docker: Lightweight Linux Containers for Consistent Development and Deployment." Linux Journal, 239.',
    '[17] Pedregosa, F., et al. (2011). "Scikit-learn: Machine Learning in Python." Journal of Machine Learning Research, 12, 2825-2830.',
    '[18] McKinney, W. (2010). "Data Structures for Statistical Computing in Python." Proceedings of the 9th Python in Science Conference, 51-56.',
    '[19] Werkzeug Contributors (2024). "Werkzeug \u2014 The Comprehensive WSGI Web Application Library." werkzeug.palletsprojects.com.',
    '[20] Hastie, T., Tibshirani, R., & Friedman, J. (2009). "The Elements of Statistical Learning." Springer.',
    '[21] James, G., Witten, D., Hastie, T., & Tibshirani, R. (2013). "An Introduction to Statistical Learning." Springer.',
    '[22] Raschka, S., & Mirjalili, V. (2019). "Python Machine Learning." Packt Publishing.',
    '[23] G\u00e9ron, A. (2019). "Hands-On Machine Learning with Scikit-Learn, Keras, and TensorFlow." O\'Reilly Media.',
    '[24] Lundberg, S. M., & Lee, S. I. (2017). "A Unified Approach to Interpreting Model Predictions." NeurIPS 2017.',
    '[25] Flask Documentation (2024). "Flask \u2014 Web Development, One Drop at a Time." flask.palletsprojects.com.',
    '[26] SQLite Documentation (2024). "SQLite \u2014 A Self-contained, Serverless, SQL Database Engine." sqlite.org.',
    '[27] Scikit-learn Documentation (2024). "Scikit-learn: Machine Learning in Python." scikit-learn.org.',
    '[28] Docker Documentation (2024). "Docker Docs \u2014 Get Started." docs.docker.com.',
    '[29] Insurance Information Institute (2024). "Facts + Statistics: Auto Insurance." iii.org.',
    '[30] Kaggle (2024). "Car Insurance Claim Prediction Datasets." kaggle.com/datasets.',
]

for ref in references:
    add_justified_text(ref, font_size=11, space_after=4)

# ============================================================
# SAVE DOCUMENT
# ============================================================
doc.save(OUTPUT_PATH)
file_size = os.path.getsize(OUTPUT_PATH) // 1024
print(f"Report saved to: {OUTPUT_PATH}")
print(f"File size: {file_size} KB")
