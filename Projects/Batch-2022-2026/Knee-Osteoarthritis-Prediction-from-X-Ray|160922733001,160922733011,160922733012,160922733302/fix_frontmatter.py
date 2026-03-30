"""Fix C2 front matter to match C16 reference format."""
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from lxml import etree

DOC_PATH = 'Knee_Osteoarthritis_Prediction_Major_Project_Report.docx'
doc = Document(DOC_PATH)
body = doc.element.body


# ============================================================
# 1. DELETE C16 LEFTOVER ABSTRACT PARAGRAPHS (80-82)
#    Para 79 already has the complete C2 abstract.
#    Paras 80-82 contain brain hemorrhage text from the C16 template.
# ============================================================
print("=== Fix 1: Remove C16 leftover abstract paragraphs ===")

# Verify the paragraphs to delete contain C16 text
paras_to_delete = []
for i in [80, 81, 82]:
    txt = doc.paragraphs[i].text.strip()
    if any(kw in txt.lower() for kw in ['brain hemorrhage', 'ct brain images', 'ct image upload']):
        paras_to_delete.append(i)
        print(f"  Para[{i}]: REMOVING (C16 leftover) — '{txt[:80]}...'")
    else:
        print(f"  Para[{i}]: KEEPING (not C16 text) — '{txt[:80]}'")

# Delete in reverse order to preserve indices
for i in sorted(paras_to_delete, reverse=True):
    elem = doc.paragraphs[i]._element
    body.remove(elem)
print(f"  Deleted {len(paras_to_delete)} C16 leftover paragraphs")


# ============================================================
# 2. ADD PAGE BREAK BETWEEN CERTIFICATE AND DECLARATION
#    C16 has a page break paragraph between them.
#    C2 is missing it — Declaration may start on same page as Certificate.
# ============================================================
print("\n=== Fix 2: Add page break before Declaration ===")

# Find "DECLARATION BY THE CANDIDATE" paragraph
decl_idx = None
for i, p in enumerate(doc.paragraphs):
    if 'DECLARATION' in p.text.upper() and 'CANDIDATE' in p.text.upper():
        decl_idx = i
        break

if decl_idx:
    # Check if there's already a page break in the paragraph before it
    prev_has_break = False
    for j in range(decl_idx - 1, max(0, decl_idx - 3), -1):
        for br in doc.paragraphs[j]._element.findall(f'.//{{{qn("w:br").split("}")[0][1:]}}}br'):
            if br.get(qn('w:type')) == 'page':
                prev_has_break = True
                break
        # Also check via the simpler path
        breaks = doc.paragraphs[j]._element.findall('.//' + qn('w:br'))
        for br in breaks:
            if br.get(qn('w:type')) == 'page':
                prev_has_break = True

    if not prev_has_break:
        # Insert a page break paragraph before the Declaration
        decl_elem = doc.paragraphs[decl_idx]._element
        pb_p = etree.SubElement(body, qn('w:p'))
        pb_r = etree.SubElement(pb_p, qn('w:r'))
        pb_br = etree.SubElement(pb_r, qn('w:br'))
        pb_br.set(qn('w:type'), 'page')
        body.remove(pb_p)
        decl_elem.addprevious(pb_p)
        print(f"  Added page break before Para[{decl_idx}] (Declaration)")
    else:
        print(f"  Page break already exists before Declaration")
else:
    print("  WARNING: Could not find Declaration paragraph")


# ============================================================
# 3. FIX TOC FRONT MATTER ENTRIES — title case + page numbers
#    C16 uses: "Abstract", "Table of Contents", "List of Figures", "List of Tables"
#    C2 has:   "ABSTRACT", "TABLE OF CONTENTS", "LIST OF FIGURES", "LIST OF TABLES"
#    Also fix page numbers to match C16 pattern.
# ============================================================
print("\n=== Fix 3: Fix TOC front matter entries ===")
toc = doc.tables[21]

# Map of corrections: (current_text_upper, new_text, new_page)
toc_fixes = {
    'SDG MAPPING': ('SDG Mapping', 'xii'),
    'ABSTRACT': ('Abstract', 'xiii'),
    'TABLE OF CONTENTS': ('Table of Contents', 'xiv'),
    'LIST OF FIGURES': ('List of Figures', 'xvi'),
    'LIST OF TABLES': ('List of Tables', 'xvii'),
}

for ri in range(len(toc.rows)):
    c0_text = toc.cell(ri, 0).text.strip()
    c0_upper = c0_text.upper()

    for key, (new_text, new_page) in toc_fixes.items():
        if c0_upper == key:
            # Fix entry text
            cell0 = toc.cell(ri, 0)
            for run in cell0.paragraphs[0].runs:
                run.text = ""
            if cell0.paragraphs[0].runs:
                cell0.paragraphs[0].runs[0].text = new_text
            else:
                r = cell0.paragraphs[0].add_run(new_text)
                r.font.name = 'Times New Roman'
                r.font.size = Pt(11)

            # Fix page number
            cell1 = toc.cell(ri, 1)
            for run in cell1.paragraphs[0].runs:
                run.text = ""
            if cell1.paragraphs[0].runs:
                cell1.paragraphs[0].runs[0].text = new_page
            else:
                r = cell1.paragraphs[0].add_run(new_page)
                r.font.name = 'Times New Roman'
                r.font.size = Pt(11)

            print(f"  Row[{ri}]: '{c0_text}' → '{new_text}' | {new_page}")
            break


# ============================================================
# SAVE
# ============================================================
doc.save(DOC_PATH)
import os
size = os.path.getsize(DOC_PATH) // 1024
print(f"\nDone. Saved: {DOC_PATH} ({size} KB)")
