"""Fix C2 Report: TOC/LOF/LOT format to match C11 reference template."""
import copy, re
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from lxml import etree

DOC_PATH = 'Knee_Osteoarthritis_Prediction_Major_Project_Report.docx'
doc = Document(DOC_PATH)


def set_cell_text(cell, text, bold=False, size=11, align=None):
    """Set cell text preserving formatting."""
    p = cell.paragraphs[0]
    # Clear existing runs
    for run in p.runs:
        run.text = ""
    if p.runs:
        r = p.runs[0]
    else:
        r = p.add_run()
    r.text = text
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.bold = bold
    # Set font in XML too
    rPr = r._element.find(qn('w:rPr'))
    if rPr is None:
        rPr = etree.SubElement(r._element, qn('w:rPr'))
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = etree.SubElement(rPr, qn('w:rFonts'))
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:cs'), 'Times New Roman')
    if align:
        pPr = p._element.find(qn('w:pPr'))
        if pPr is None:
            pPr = etree.SubElement(p._element, qn('w:pPr'))
            p._element.insert(0, pPr)
        jc = pPr.find(qn('w:jc'))
        if jc is None:
            jc = etree.SubElement(pPr, qn('w:jc'))
        jc.set(qn('w:val'), align)


def insert_row_before(table, ref_row_idx):
    """Insert a new row before ref_row_idx by copying the first data row."""
    tbl_elem = table._tbl
    ref_tr = table.rows[ref_row_idx]._tr
    new_tr = copy.deepcopy(ref_tr)
    # Clear all text in the new row
    for t_elem in new_tr.iter(qn('w:t')):
        t_elem.text = ""
    ref_tr.addprevious(new_tr)
    return new_tr


def remove_empty_trailing_rows(table):
    """Remove empty rows at the end of the table."""
    removed = 0
    while len(table.rows) > 1:
        last = table.rows[-1]
        if all(c.text.strip() == '' for c in last.cells):
            table._tbl.remove(last._tr)
            removed += 1
        else:
            break
    return removed


# ============================================================
# 1. FIX LIST OF FIGURES (Table 22)
#    Current: "Fig 1.1: Title | Title | 4" (no header)
#    Target:  Header: "Fig. No. | Title | Page No."
#             Data:   "Fig. 1.1 | Title | 4"
# ============================================================
print("=== Fixing List of Figures Format ===")
lof = doc.tables[22]

# Remove trailing empty rows first
rem = remove_empty_trailing_rows(lof)
print(f"  Removed {rem} empty trailing rows")

# Fix data rows: col0 "Fig 1.1: Title" → "Fig. 1.1"
for ri in range(len(lof.rows)):
    c0_text = lof.cell(ri, 0).text.strip()
    # Extract "Fig X.X" from "Fig X.X: Title"
    m = re.match(r'(Fig)\s*(\d+\.\d+)', c0_text)
    if m:
        fig_no = f"Fig. {m.group(2)}"
        set_cell_text(lof.cell(ri, 0), fig_no, size=11)

# Insert header row at position 0
header_tr = insert_row_before(lof, 0)
# Now lof.rows[0] is the new header - need to re-read table
# Set header text using XML directly
cells = header_tr.findall(qn('w:tc'))
headers = ["Fig. No.", "Title", "Page No."]
for ci, htext in enumerate(headers):
    if ci < len(cells):
        t_elem = cells[ci].find('.//' + qn('w:t'))
        if t_elem is not None:
            t_elem.text = htext
        # Bold the header
        r_elem = cells[ci].find('.//' + qn('w:r'))
        if r_elem is not None:
            rPr = r_elem.find(qn('w:rPr'))
            if rPr is None:
                rPr = etree.SubElement(r_elem, qn('w:rPr'))
            b = rPr.find(qn('w:b'))
            if b is None:
                etree.SubElement(rPr, qn('w:b'))

# Mark header row to repeat on each page
trPr = header_tr.find(qn('w:trPr'))
if trPr is None:
    trPr = etree.SubElement(header_tr, qn('w:trPr'))
    header_tr.insert(0, trPr)
tblHeader = trPr.find(qn('w:tblHeader'))
if tblHeader is None:
    etree.SubElement(trPr, qn('w:tblHeader'))

print(f"  Added header row: Fig. No. | Title | Page No.")
print(f"  Fixed {len(lof.rows)-1} figure number formats (Fig. X.X)")


# ============================================================
# 2. FIX LIST OF TABLES (Table 23)
#    Current: "Table 2.1: Title | Title | 8" (no header)
#    Target:  Header: "Table No. | Title | Page No."
#             Data:   "Table 2.1 | Title | 8"
# ============================================================
print("\n=== Fixing List of Tables Format ===")
lot = doc.tables[23]

# Remove trailing empty rows
rem = remove_empty_trailing_rows(lot)
print(f"  Removed {rem} empty trailing rows")

# Fix data rows: col0 "Table 2.1: Title" → "Table 2.1"
for ri in range(len(lot.rows)):
    c0_text = lot.cell(ri, 0).text.strip()
    m = re.match(r'(Table)\s*(\d+\.\d+)', c0_text)
    if m:
        tbl_no = f"Table {m.group(2)}"
        set_cell_text(lot.cell(ri, 0), tbl_no, size=11)

# Insert header row
header_tr = insert_row_before(lot, 0)
cells = header_tr.findall(qn('w:tc'))
headers = ["Table No.", "Title", "Page No."]
for ci, htext in enumerate(headers):
    if ci < len(cells):
        t_elem = cells[ci].find('.//' + qn('w:t'))
        if t_elem is not None:
            t_elem.text = htext
        r_elem = cells[ci].find('.//' + qn('w:r'))
        if r_elem is not None:
            rPr = r_elem.find(qn('w:rPr'))
            if rPr is None:
                rPr = etree.SubElement(r_elem, qn('w:rPr'))
            b = rPr.find(qn('w:b'))
            if b is None:
                etree.SubElement(rPr, qn('w:b'))

trPr = header_tr.find(qn('w:trPr'))
if trPr is None:
    trPr = etree.SubElement(header_tr, qn('w:trPr'))
    header_tr.insert(0, trPr)
tblHeader = trPr.find(qn('w:tblHeader'))
if tblHeader is None:
    etree.SubElement(trPr, qn('w:tblHeader'))

print(f"  Added header row: Table No. | Title | Page No.")
print(f"  Fixed {len(lot.rows)-1} table number formats")


# ============================================================
# 3. FIX TABLE OF CONTENTS (Table 21)
#    Current: starts with "ABSTRACT | i" (missing preliminary pages)
#    Target:  starts with "Title Page | i", "Certificate | ii", etc.
#    Reference: C11 format
# ============================================================
print("\n=== Fixing Table of Contents ===")
toc = doc.tables[21]

# Remove empty trailing rows
rem = remove_empty_trailing_rows(toc)
print(f"  Removed {rem} empty trailing rows")

# Preliminary pages to add at top (before ABSTRACT)
prelim_pages = [
    ("Title Page", "i"),
    ("Certificate", "ii"),
    ("Declaration", "iii"),
    ("Acknowledgment", "iv"),
    ("Vision & Mission of the Institute", "v"),
    ("Vision & Mission of the Department", "vi"),
    ("Program Educational Objectives (PEOs)", "vii"),
    ("Program Outcomes (POs)", "viii"),
    ("Program Specific Outcomes (PSOs)", "ix"),
    ("Course Outcomes", "x"),
    ("Course Articulation Matrix", "xi"),
    ("SDG Mapping", "xi"),
]

# Check if preliminary pages already exist
first_entry = toc.cell(0, 0).text.strip()
if first_entry.upper().startswith("ABSTRACT") or first_entry.upper().startswith("TABLE OF CONTENTS"):
    # Need to add preliminary pages
    # Insert rows before the ABSTRACT row (forward order, each before ABSTRACT)
    abstract_tr = toc.rows[0]._tr
    for title, page_no in prelim_pages:
        new_tr = copy.deepcopy(abstract_tr)
        cells = new_tr.findall(qn('w:tc'))
        if len(cells) >= 2:
            for t_elem in cells[0].iter(qn('w:t')):
                t_elem.text = ""
            first_t = cells[0].find('.//' + qn('w:t'))
            if first_t is not None:
                first_t.text = title
            for t_elem in cells[1].iter(qn('w:t')):
                t_elem.text = ""
            first_t2 = cells[1].find('.//' + qn('w:t'))
            if first_t2 is not None:
                first_t2.text = page_no
        abstract_tr.addprevious(new_tr)
    print(f"  Added {len(prelim_pages)} preliminary page entries")

    # Update existing front matter page numbers
    # After adding 12 prelim entries, re-map page numbers
    page_map = {
        "ABSTRACT": "xii",
        "TABLE OF CONTENTS": "xiii",
        "LIST OF FIGURES": "xv",
        "LIST OF TABLES": "xvi",
    }
    for ri in range(len(prelim_pages), len(toc.rows)):
        c0 = toc.cell(ri, 0).text.strip().upper()
        for key, new_page in page_map.items():
            if c0 == key:
                set_cell_text(toc.cell(ri, 1), new_page, size=11)
                print(f"  Updated '{key}' page → {new_page}")
                break
else:
    print(f"  Preliminary pages already present (first entry: '{first_entry}')")


# ============================================================
# 4. KEEP TABLES ON ONE PAGE (for body tables, not TOC/LOF/LOT)
#    Apply cantSplit + keepNext only to tables that fit on one page
#    Skip large tables (>20 rows) like TOC, LOF, LOT
# ============================================================
print("\n=== Keep Tables on One Page ===")
table_count = 0
for ti, tbl in enumerate(doc.tables):
    # Skip very large tables (TOC=21, LOF=22, LOT=23 and any >25 rows)
    if len(tbl.rows) > 25:
        continue

    tbl_elem = tbl._tbl
    for row in tbl.rows:
        tr = row._tr
        trPr = tr.find(qn('w:trPr'))
        if trPr is None:
            trPr = etree.SubElement(tr, qn('w:trPr'))
            tr.insert(0, trPr)
        cant_split = trPr.find(qn('w:cantSplit'))
        if cant_split is None:
            etree.SubElement(trPr, qn('w:cantSplit'))

        for cell in row.cells:
            for p in cell.paragraphs:
                pPr = p._element.find(qn('w:pPr'))
                if pPr is None:
                    pPr = etree.SubElement(p._element, qn('w:pPr'))
                    p._element.insert(0, pPr)
                if pPr.find(qn('w:keepNext')) is None:
                    etree.SubElement(pPr, qn('w:keepNext'))
                if pPr.find(qn('w:keepLines')) is None:
                    etree.SubElement(pPr, qn('w:keepLines'))

    table_count += 1

# For large tables (TOC, LOF, LOT), only apply cantSplit (no keepNext)
# Remove keepNext from these tables if previously applied
for ti in [21, 22, 23]:
    if ti < len(doc.tables):
        tbl = doc.tables[ti]
        for row in tbl.rows:
            tr = row._tr
            trPr = tr.find(qn('w:trPr'))
            if trPr is None:
                trPr = etree.SubElement(tr, qn('w:trPr'))
                tr.insert(0, trPr)
            # Keep cantSplit (don't split a row across pages)
            if trPr.find(qn('w:cantSplit')) is None:
                etree.SubElement(trPr, qn('w:cantSplit'))
            # Remove keepNext from large tables (it was causing issues)
            for cell in row.cells:
                for p in cell.paragraphs:
                    pPr = p._element.find(qn('w:pPr'))
                    if pPr is not None:
                        kn = pPr.find(qn('w:keepNext'))
                        if kn is not None:
                            pPr.remove(kn)

print(f"  Applied keep-together to {table_count} body tables")
print(f"  Cleaned keepNext from TOC/LOF/LOT tables")


# ============================================================
# SAVE
# ============================================================
doc.save(DOC_PATH)
import os
size = os.path.getsize(DOC_PATH) // 1024
print(f"\nDone. Saved: {DOC_PATH} ({size} KB)")
