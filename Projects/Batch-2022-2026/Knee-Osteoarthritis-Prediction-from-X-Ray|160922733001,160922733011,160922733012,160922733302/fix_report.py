"""Fix C2 Knee OA Report: LOF/LOT descriptions, literature review, figure placement, table keep-together."""
import copy
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from lxml import etree

DOC_PATH = 'Knee_Osteoarthritis_Prediction_Major_Project_Report.docx'
doc = Document(DOC_PATH)

# ============================================================
# 1. FIX LIST OF FIGURES (Table 22) - col1 descriptions
# ============================================================
print("=== Fixing List of Figures ===")
lof_correct = [
    "Traditional vs AI-Based KOA Diagnosis",
    "System Architecture Diagram",
    "Use Case Diagram",
    "Class Diagram",
    "Sequence Diagram",
    "Activity Diagram",
    "UI Wireframe — Prediction Interface",
    "ER / Database Schema Diagram",
    "Development Phase Diagram",
    "Login Page",
    "Registration Page",
    "Home Dashboard with Stats",
    "X-Ray Upload Interface",
    "Prediction Result — Normal (Grade 0)",
    "Prediction Result — Severe (Grade 4)",
    "Prediction History",
    "Model Comparison Dashboard",
    "Per-Class Accuracy Chart",
    "About Page",
    "Sample X-Ray — Normal",
    "Sample X-Ray — Doubtful",
    "Sample X-Ray — Mild",
    "Sample X-Ray — Moderate",
    "Sample X-Ray — Severe",
    "Model Accuracy Comparison",
    "F1 Score Comparison",
    "Prediction Distribution Chart",
    "Mobile Responsive View",
]

tbl_lof = doc.tables[22]
for ri in range(min(len(lof_correct), len(tbl_lof.rows))):
    cell = tbl_lof.cell(ri, 1)
    for p in cell.paragraphs:
        for run in p.runs:
            run.text = ""
        if p.runs:
            p.runs[0].text = lof_correct[ri]
        else:
            run = p.add_run(lof_correct[ri])
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'
print(f"  Fixed {len(lof_correct)} figure descriptions")

# ============================================================
# 2. FIX LIST OF TABLES (Table 23) - col1 descriptions
# ============================================================
print("=== Fixing List of Tables ===")
lot_correct = [
    "Literature Survey Comparison",
    "Feasibility Study",
    "System Feature Requirements",
    "Non-Functional Requirements",
    "Hardware Requirements",
    "Software Requirements",
    "Users Table Schema",
    "Predictions Table Schema",
    "Module Description",
    "Test Cases — Authentication",
    "Test Cases — Upload & Prediction",
    "Test Cases — History & Dashboard",
    "Test Cases — Admin & Security",
    "Model Performance Comparison",
    "Per-Class Accuracy — MobileNetV2",
]

tbl_lot = doc.tables[23]
for ri in range(min(len(lot_correct), len(tbl_lot.rows))):
    cell = tbl_lot.cell(ri, 1)
    for p in cell.paragraphs:
        for run in p.runs:
            run.text = ""
        if p.runs:
            p.runs[0].text = lot_correct[ri]
        else:
            run = p.add_run(lot_correct[ri])
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'
print(f"  Fixed {len(lot_correct)} table descriptions")

# ============================================================
# 3. EXPAND LITERATURE REVIEW (add 2.7-2.15 before summary)
# ============================================================
print("=== Expanding Literature Review ===")

new_lit_sections = [
    ("2.7  Data Augmentation for Medical Imaging (Shorten & Khoshgoftaar, 2019)",
     "Shorten and Khoshgoftaar conducted a comprehensive survey on data augmentation techniques for deep learning in medical imaging. "
     "They demonstrated that geometric transformations (rotation, flipping, scaling), color-space adjustments, and elastic deformations "
     "significantly improve model generalization when training data is limited. For knee X-ray classification, augmentation addresses "
     "the challenge of class imbalance across KL grades, where Grade 0 (Normal) typically dominates. Their findings show that augmentation "
     "can improve accuracy by 5-15% on small medical datasets. Our system applies random rotation (±15°), horizontal flip, brightness "
     "adjustment, and contrast variation during MobileNetV2 training to prevent overfitting on the knee X-ray dataset."),

    ("2.8  MobileNetV2: Inverted Residuals and Linear Bottlenecks (Sandler et al., 2018)",
     "Sandler et al. introduced MobileNetV2 with inverted residual blocks and linear bottleneck layers, achieving ImageNet accuracy "
     "comparable to larger models while requiring 10x fewer parameters. The architecture uses depthwise separable convolutions to "
     "reduce computational cost from O(D²·M·N·K²) to O(D²·M·(K²+N)), where D is spatial dimension, M and N are channel counts, "
     "and K is kernel size. The inverted residual structure expands channels in the middle layer and compresses at the output, "
     "preserving information flow through shortcut connections. For our knee OA system, MobileNetV2's lightweight architecture "
     "enables real-time inference in the Flask web application without GPU requirements, making deployment feasible on standard hardware."),

    ("2.9  Flask Web Framework for ML Deployment (Grinberg, 2018)",
     "Grinberg provided a comprehensive guide to Flask web development, covering routing, templates, database integration, and "
     "deployment patterns. Flask's microframework architecture provides minimal overhead while supporting extensions for database "
     "ORM, authentication, and file upload handling. The Jinja2 template engine enables server-side HTML rendering with dynamic "
     "content injection. For ML deployment, Flask's lightweight nature makes it ideal for serving pre-trained models through REST "
     "endpoints. Our system uses Flask to serve the MobileNetV2 model, handle X-ray image uploads, manage user sessions, store "
     "predictions in SQLite, and render interactive dashboards with Chart.js visualizations."),

    ("2.10  SQLite for Embedded Database Applications (Owens, 2006)",
     "Owens described SQLite as a self-contained, serverless, zero-configuration relational database engine suitable for embedded "
     "applications. SQLite stores the entire database in a single cross-platform file, eliminating the need for a separate database "
     "server process. ACID transaction support ensures data integrity even during system failures. For our knee OA prediction platform, "
     "SQLite (knee_osteoarthritis.db) stores user authentication data (hashed passwords) and prediction history (image paths, KL grades, "
     "confidence scores, timestamps) with foreign key relationships. The serverless architecture simplifies deployment and eliminates "
     "database administration overhead."),

    ("2.11  Kellgren-Lawrence Grading System for Knee OA (Kellgren & Lawrence, 1957)",
     "Kellgren and Lawrence established the radiographic grading system that remains the gold standard for assessing knee osteoarthritis "
     "severity. The 5-grade scale classifies knee X-rays as: Grade 0 (Normal — no radiographic features of OA), Grade 1 (Doubtful — "
     "minute osteophytes of doubtful significance), Grade 2 (Mild — definite osteophytes with possible joint space narrowing), Grade 3 "
     "(Moderate — moderate osteophytes, definite narrowing, some sclerosis), and Grade 4 (Severe — large osteophytes, marked narrowing, "
     "severe sclerosis, bone deformity). Our system automates this grading using deep learning, enabling consistent and rapid assessment "
     "that reduces inter-observer variability inherent in manual radiographic interpretation."),

    ("2.12  ImageNet and Transfer Learning (Deng et al., 2009)",
     "Deng et al. created ImageNet, a large-scale hierarchical image database containing over 14 million labeled images across 20,000+ "
     "categories. ImageNet pre-training provides neural networks with general visual feature extraction capabilities (edges, textures, "
     "shapes) that transfer effectively to domain-specific tasks through fine-tuning. For medical imaging, transfer learning from ImageNet "
     "overcomes the limitation of small labeled datasets by leveraging features learned from millions of natural images. Our MobileNetV2 "
     "model uses ImageNet pre-trained weights as initialization, then fine-tunes all layers on knee X-ray images to specialize the "
     "feature representations for KL grade classification."),

    ("2.13  Classification Metrics for Medical AI (Powers, 2011)",
     "Powers provided a comprehensive analysis of evaluation metrics for classification systems, distinguishing between accuracy, "
     "precision, recall, F1-score, and AUC-ROC. In medical AI applications, these metrics have clinical significance: precision measures "
     "the proportion of positive predictions that are correct (minimizing false alarms), recall measures the proportion of actual "
     "positives detected (minimizing missed diagnoses), and F1-score balances both. For multi-class KL grading, per-class accuracy "
     "reveals whether the model performs uniformly across all 5 grades or favors certain classes. Our system evaluates MobileNetV2 "
     "using accuracy, precision, recall, F1-score, and per-class accuracy, providing comprehensive performance assessment."),

    ("2.14  Bootstrap 5 for Responsive Web Design (Bootstrap, 2021)",
     "Bootstrap 5 removed the jQuery dependency and introduced improved utility classes, responsive grid system, and dark theme support. "
     "The mobile-first approach ensures that web applications render correctly across devices from smartphones to desktop monitors. "
     "For medical applications, responsive design enables clinicians to access the system on tablets during patient consultations or "
     "on large monitors in radiology departments. Our knee OA platform uses Bootstrap 5 with a custom dark gradient theme, card "
     "components for organized data display, drag-and-drop file upload zones, and Chart.js integration for interactive model "
     "performance dashboards."),

    ("2.15  Docker Containerization for ML Applications (Merkel, 2014)",
     "Merkel described Docker as a platform for creating lightweight, portable containers that package applications with all "
     "dependencies. Docker eliminates environment inconsistencies by ensuring identical development, testing, and production "
     "configurations. For ML applications, Docker containers bundle Python, PyTorch, Flask, and all libraries into a reproducible "
     "image that runs consistently on any host. Our system includes a Dockerfile based on Python 3.11-slim that installs all "
     "dependencies (Flask, PyTorch, scikit-learn, Pillow), copies model files, and configures the Flask application to run on "
     "port 5011, enabling one-command deployment on any Docker-capable machine."),
]

# Find the paragraph index of "2.7  Literature Review Summary"
summary_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip().startswith("2.7") and "Summary" in p.text:
        summary_idx = i
        break

if summary_idx:
    print(f"  Found summary at Para[{summary_idx}]")
    body = doc.element.body
    anchor = doc.paragraphs[summary_idx]._element

    # FIRST: Renumber old "2.7  Summary" → "2.16  Summary" BEFORE inserting new sections
    for run in doc.paragraphs[summary_idx].runs:
        if "2.7" in run.text:
            run.text = run.text.replace("2.7", "2.16")
            break
    for elem in anchor.iter(qn('w:t')):
        if elem.text and "2.7" in elem.text:
            elem.text = elem.text.replace("2.7", "2.16")

    # Insert new sections BEFORE the summary in correct order (2.7, 2.8, ... 2.15)
    for heading_text, content_text in new_lit_sections:
        # Create heading paragraph
        heading_p = etree.SubElement(body, qn('w:p'))
        heading_pPr = etree.SubElement(heading_p, qn('w:pPr'))
        heading_sp = etree.SubElement(heading_pPr, qn('w:spacing'))
        heading_sp.set(qn('w:before'), '240')
        heading_sp.set(qn('w:after'), '120')
        heading_r = etree.SubElement(heading_p, qn('w:r'))
        heading_rPr = etree.SubElement(heading_r, qn('w:rPr'))
        etree.SubElement(heading_rPr, qn('w:b'))
        heading_sz = etree.SubElement(heading_rPr, qn('w:sz'))
        heading_sz.set(qn('w:val'), '24')
        heading_szCs = etree.SubElement(heading_rPr, qn('w:szCs'))
        heading_szCs.set(qn('w:val'), '24')
        heading_font = etree.SubElement(heading_rPr, qn('w:rFonts'))
        heading_font.set(qn('w:ascii'), 'Times New Roman')
        heading_font.set(qn('w:hAnsi'), 'Times New Roman')
        heading_t = etree.SubElement(heading_r, qn('w:t'))
        heading_t.text = heading_text
        heading_t.set(qn('xml:space'), 'preserve')
        body.remove(heading_p)
        anchor.addprevious(heading_p)

        # Create content paragraph
        content_p = etree.SubElement(body, qn('w:p'))
        content_pPr = etree.SubElement(content_p, qn('w:pPr'))
        content_jc = etree.SubElement(content_pPr, qn('w:jc'))
        content_jc.set(qn('w:val'), 'both')
        content_ind = etree.SubElement(content_pPr, qn('w:ind'))
        content_ind.set(qn('w:firstLine'), '720')
        content_sp = etree.SubElement(content_pPr, qn('w:spacing'))
        content_sp.set(qn('w:after'), '120')
        content_r = etree.SubElement(content_p, qn('w:r'))
        content_rPr = etree.SubElement(content_r, qn('w:rPr'))
        content_sz = etree.SubElement(content_rPr, qn('w:sz'))
        content_sz.set(qn('w:val'), '24')
        content_szCs = etree.SubElement(content_rPr, qn('w:szCs'))
        content_szCs.set(qn('w:val'), '24')
        content_font = etree.SubElement(content_rPr, qn('w:rFonts'))
        content_font.set(qn('w:ascii'), 'Times New Roman')
        content_font.set(qn('w:hAnsi'), 'Times New Roman')
        content_t = etree.SubElement(content_r, qn('w:t'))
        content_t.text = content_text
        content_t.set(qn('xml:space'), 'preserve')
        body.remove(content_p)
        anchor.addprevious(content_p)

    print(f"  Added {len(new_lit_sections)} new literature sections (2.7-2.15)")
    print(f"  Renumbered old 2.7 Summary → 2.16")
else:
    print("  WARNING: Could not find literature summary paragraph!")

# ============================================================
# 4. UPDATE TOC (Table 21) - add new lit review entries
# ============================================================
print("=== Updating Table of Contents ===")
tbl_toc = doc.tables[21]

# Find the row with "2.7  Literature Review Summary"
toc_summary_row = None
for ri in range(len(tbl_toc.rows)):
    c0 = tbl_toc.cell(ri, 0).text.strip()
    if "2.7" in c0 and "Summary" in c0:
        toc_summary_row = ri
        break

if toc_summary_row:
    # Update existing summary row to 2.16
    cell0 = tbl_toc.cell(toc_summary_row, 0)
    for run in cell0.paragraphs[0].runs:
        if "2.7" in run.text:
            run.text = run.text.replace("2.7", "2.16")
    for elem in cell0._tc.iter(qn('w:t')):
        if elem.text and "2.7" in elem.text:
            elem.text = elem.text.replace("2.7", "2.16")

    # Insert new rows before summary row
    new_toc_entries = [
        ("2.7  Data Augmentation for Medical Imaging", ""),
        ("2.8  MobileNetV2: Inverted Residuals", ""),
        ("2.9  Flask Web Framework for ML Deployment", ""),
        ("2.10  SQLite for Embedded Database Applications", ""),
        ("2.11  Kellgren-Lawrence Grading System", ""),
        ("2.12  ImageNet and Transfer Learning", ""),
        ("2.13  Classification Metrics for Medical AI", ""),
        ("2.14  Bootstrap 5 for Responsive Web Design", ""),
        ("2.15  Docker Containerization for ML Applications", ""),
    ]

    # Add rows to the TOC table
    tbl_elem = tbl_toc._tbl
    summary_tr = tbl_toc.rows[toc_summary_row]._tr

    for entry_text, page in new_toc_entries:
        # Create a new table row by copying the summary row structure
        new_tr = copy.deepcopy(summary_tr)
        # Clear and set text
        cells = new_tr.findall(qn('w:tc'))
        if len(cells) >= 2:
            # Cell 0 - entry text
            for t_elem in cells[0].iter(qn('w:t')):
                t_elem.text = ""
            first_t = cells[0].find('.//' + qn('w:t'))
            if first_t is not None:
                first_t.text = entry_text
            # Cell 1 - page number
            for t_elem in cells[1].iter(qn('w:t')):
                t_elem.text = ""
            first_t2 = cells[1].find('.//' + qn('w:t'))
            if first_t2 is not None:
                first_t2.text = page

        summary_tr.addprevious(new_tr)

    print(f"  Added {len(new_toc_entries)} new TOC entries")
    print(f"  Updated 2.7 Summary → 2.16 in TOC")
else:
    print("  WARNING: Could not find TOC summary row!")

# ============================================================
# 5. FIX FIGURE PLACEMENT IN CHAPTER 4 - move images after content
# ============================================================
print("=== Fixing Figure Placement ===")

# The issue: in Ch4, images appear right after heading, before description
# We need to move image paragraphs to AFTER the description text
# Target: paras with images in Ch4 range (210-234)

body = doc.element.body
fixed_count = 0

# Get all paragraph elements
all_paras = list(body.iterchildren(qn('w:p')))

for i, p_elem in enumerate(all_paras):
    # Check if this paragraph has an image
    drawings = p_elem.findall('.//' + qn('w:drawing'))
    if not drawings:
        continue

    # Get paragraph index in doc.paragraphs
    try:
        p_idx = list(body.iterchildren(qn('w:p'))).index(p_elem)
    except:
        continue

    # Only fix Ch4 figures (skip title page logo, Ch7 screenshots)
    # Ch4 figures are roughly between "CHAPTER 4" and "CHAPTER 5"
    # Check surrounding text
    if i + 1 < len(all_paras):
        next_text = ""
        for t in all_paras[i + 1].iter(qn('w:t')):
            if t.text:
                next_text += t.text
        next_text = next_text.strip()

        prev_text = ""
        if i > 0:
            for t in all_paras[i - 1].iter(qn('w:t')):
                if t.text:
                    prev_text += t.text
            prev_text = prev_text.strip()

        # Only move if: heading is before image AND description is after
        # Pattern: heading → IMAGE → description text
        if (prev_text.startswith(("4.3.1", "4.3.2", "4.3.3", "4.3.4", "4.5.1")) and
                not next_text.startswith(("4.", "CHAPTER", "Fig"))):
            # Find the next paragraph that starts a new section or is empty
            # Move image after description text
            j = i + 1
            insert_after = None
            while j < len(all_paras):
                jtext = ""
                for t in all_paras[j].iter(qn('w:t')):
                    if t.text:
                        jtext += t.text
                jtext = jtext.strip()
                # Stop at next section heading or another image
                if (jtext.startswith(("4.", "CHAPTER", "5.")) or
                        all_paras[j].findall('.//' + qn('w:drawing'))):
                    insert_after = all_paras[j - 1]
                    break
                j += 1

            if insert_after is not None and insert_after is not p_elem:
                body.remove(p_elem)
                insert_after.addnext(p_elem)
                fixed_count += 1
                print(f"  Moved figure after: {prev_text[:50]}...")

    # Re-read all_paras since we modified the tree
    all_paras = list(body.iterchildren(qn('w:p')))

print(f"  Fixed {fixed_count} figure placements")

# ============================================================
# 6. KEEP TABLES ON ONE PAGE
# ============================================================
print("=== Adding Keep-Table-On-One-Page ===")
table_count = 0
for tbl in doc.tables:
    tbl_elem = tbl._tbl
    tblPr = tbl_elem.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = etree.SubElement(tbl_elem, qn('w:tblPr'))
        tbl_elem.insert(0, tblPr)

    for row in tbl.rows:
        tr = row._tr
        trPr = tr.find(qn('w:trPr'))
        if trPr is None:
            trPr = etree.SubElement(tr, qn('w:trPr'))
            tr.insert(0, trPr)
        # cantSplit keeps the row from splitting across pages
        cant_split = trPr.find(qn('w:cantSplit'))
        if cant_split is None:
            etree.SubElement(trPr, qn('w:cantSplit'))

        # Also set keepNext on each cell's paragraphs (except last row)
        for cell in row.cells:
            for p in cell.paragraphs:
                pPr = p._element.find(qn('w:pPr'))
                if pPr is None:
                    pPr = etree.SubElement(p._element, qn('w:pPr'))
                    p._element.insert(0, pPr)
                keep_next = pPr.find(qn('w:keepNext'))
                if keep_next is None:
                    etree.SubElement(pPr, qn('w:keepNext'))
                keep_lines = pPr.find(qn('w:keepLines'))
                if keep_lines is None:
                    etree.SubElement(pPr, qn('w:keepLines'))

    # Compress cell margins for better fit
    for row in tbl.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is None:
                tcPr = etree.SubElement(tc, qn('w:tcPr'))
            tcMar = tcPr.find(qn('w:tcMar'))
            if tcMar is None:
                tcMar = etree.SubElement(tcPr, qn('w:tcMar'))
            for side in ['top', 'bottom']:
                elem = tcMar.find(qn(f'w:{side}'))
                if elem is None:
                    elem = etree.SubElement(tcMar, qn(f'w:{side}'))
                elem.set(qn('w:w'), '30')
                elem.set(qn('w:type'), 'dxa')

    table_count += 1

print(f"  Applied keep-together to {table_count} tables")

# ============================================================
# SAVE
# ============================================================
doc.save(DOC_PATH)
import os
size = os.path.getsize(DOC_PATH) // 1024
print(f"\nDone. Saved: {DOC_PATH} ({size} KB)")
