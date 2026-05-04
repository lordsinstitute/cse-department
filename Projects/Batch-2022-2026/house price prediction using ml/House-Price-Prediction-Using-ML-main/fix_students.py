"""Replace student names/rolls in C4 House Price Prediction Report."""

from docx import Document

DOC_PATH = '/Users/shoukathali/lord-major-projects/IV-C Projects/C4/House_Price_Prediction_Using_ML_Major_Project_Report.docx'

NEW_STUDENTS = [
    ('Mukannaf Ahmed', '160922733133'),
    ('Syed Muhtishim Najeed', '160922733154'),
    ('Md. Faizan', '160922733174'),
]

OLD_STUDENTS = [
    ('Muhammad Aasim Uz Zaman', '160922733020'),
    ('Syed Altamash Uddin Siddiqui', '160922733032'),
    ('Nawaz Khan', '160922733037'),
    ('Faiz Ur Rahman', '160922733049'),
]

doc = Document(DOC_PATH)

# ── 1. Replace students in tables 0, 5, 7 ──
for ti in [0, 5, 7]:
    tbl = doc.tables[ti]
    print(f'\nTable {ti}:')
    for ri in range(min(3, len(tbl.rows))):
        name, roll = NEW_STUDENTS[ri]
        for run in tbl.rows[ri].cells[0].paragraphs[0].runs:
            run.text = ''
        tbl.rows[ri].cells[0].paragraphs[0].runs[0].text = name
        for run in tbl.rows[ri].cells[1].paragraphs[0].runs:
            run.text = ''
        tbl.rows[ri].cells[1].paragraphs[0].runs[0].text = roll
        print(f'  Row {ri}: {name} | {roll}')

    # Remove 4th row if it has old student (not the Lords Institute row)
    if len(tbl.rows) >= 4:
        cell_text = tbl.rows[3].cells[0].text.strip()
        if 'Lords' not in cell_text and cell_text:
            tbl.rows[3]._tr.getparent().remove(tbl.rows[3]._tr)
            print(f'  Row 3: Removed (was "{cell_text}")')
        else:
            print(f'  Row 3: Kept ("{cell_text}")')

# ── 2. Replace students in certificate paragraph ──
old_names = 'Muhammad Aasim Uz Zaman (160922733020), Syed Altamash Uddin Siddiqui (160922733032), Nawaz Khan (160922733037), and Faiz Ur Rahman (160922733049)'
new_names = 'Mukannaf Ahmed (160922733133), Syed Muhtishim Najeed (160922733154), and Md. Faizan (160922733174)'

for i, para in enumerate(doc.paragraphs):
    if old_names in para.text:
        full = para.text.replace(old_names, new_names)
        for run in para.runs:
            run.text = ''
        para.runs[0].text = full
        print(f'\n[{i}] Certificate: replaced student names')
        break

doc.save(DOC_PATH)
print(f'\nDocument saved: {DOC_PATH}')
