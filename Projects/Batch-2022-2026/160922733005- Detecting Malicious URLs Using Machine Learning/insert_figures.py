"""Insert all figures into C12 Malicious URL Detection Report."""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from lxml import etree
from docx.text.paragraph import Paragraph
import os

DOC_PATH = '/Users/shoukathali/lord-major-projects/IV-C Projects/C12/Detecting_Malicious_URLs_Using_ML_Major_Project_Report.docx'
FIGURES_DIR = os.path.join(os.path.dirname(__file__), 'figures')

# Ch4-5 placeholder figures
CH4_MAP = {
    'Fig 4.1: System Architecture Diagram': ('fig_4_1_architecture.png', 5.5),
    'Fig 4.2: Use Case Diagram': ('fig_4_2_usecase.png', 5.5),
    'Fig 4.3: Class Diagram': ('fig_4_3_class.png', 5.5),
    'Fig 4.4: Sequence Diagram': ('fig_4_4_sequence.png', 5.5),
    'Fig 4.5: Activity Diagram': ('fig_4_5_activity.png', 4.5),
    'Fig 4.6: UI Wireframe': ('fig_4_6_wireframe.png', 5.5),
    'Fig 4.7: ER / Database Schema Diagram': ('fig_4_7_er_diagram.png', 5.5),
    'Fig 5.1: Development Phase Diagram': ('fig_5_1_phases.png', 5.0),
}

# Ch7 screenshot figures — use colon suffix to avoid substring matching
CH7_MAP = {
    'Fig 7.1:': ('fig_7_1_login.png', 5.0),
    'Fig 7.2:': ('fig_7_2_register.png', 5.0),
    'Fig 7.3:': ('fig_7_3_home.png', 5.0),
    'Fig 7.4:': ('fig_7_4_predict.png', 5.5),
    'Fig 7.5:': ('fig_7_5_result_legit.png', 5.5),
    'Fig 7.6:': ('fig_7_6_result_malicious.png', 5.5),
    'Fig 7.7:': ('fig_7_7_history.png', 5.5),
    'Fig 7.8:': ('fig_7_8_visualize.png', 5.5),
    'Fig 7.9:': ('fig_7_9_dashboard.png', 5.5),
    'Fig 7.10:': ('fig_7_10_feature_importance.png', 5.0),
    'Fig 7.11:': ('fig_7_11_confusion.png', 5.0),
    'Fig 7.12:': ('fig_7_12_about.png', 5.0),
    'Fig 7.13:': ('fig_7_13_admin.png', 5.5),
    'Fig 7.14:': ('fig_7_14_mobile.png', 3.0),
}

doc = Document(DOC_PATH)
inserted = 0

# Insert Ch4-5 placeholder figures
print("Inserting Chapter 4-5 figures...\n")
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if not text.startswith('['):
        continue
    for key, (img_file, width) in CH4_MAP.items():
        if key in text and 'to be inserted' in text:
            img_path = os.path.join(FIGURES_DIR, img_file)
            if not os.path.exists(img_path):
                print(f"  WARNING: {img_path} not found"); continue
            for run in para.runs: run.text = ''
            run = para.add_run()
            run.add_picture(img_path, width=Inches(width))
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            inserted += 1
            print(f"  [{i}] Inserted: {img_file}")
            break

# Insert Ch7 screenshot figures
print("\nInserting Chapter 7 figures...\n")
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if not text.startswith('[Fig 7.'):
        continue
    for key, (img_file, width) in CH7_MAP.items():
        if key in text and 'to be inserted' in text:
            img_path = os.path.join(FIGURES_DIR, img_file)
            if not os.path.exists(img_path):
                print(f"  WARNING: {img_path} not found"); continue
            for run in para.runs: run.text = ''
            run = para.add_run()
            run.add_picture(img_path, width=Inches(width))
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            inserted += 1
            print(f"  [{i}] Inserted: {img_file}")
            break

# Insert Fig 1.1 before "1.5  Existing System"
fig_1_1_done = False
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if '1.5' in text and 'Existing System' in text and not fig_1_1_done:
        img_path = os.path.join(FIGURES_DIR, 'fig_1_1_comparison.png')
        if os.path.exists(img_path):
            new_elem = etree.SubElement(para._element.getparent(), qn('w:p'))
            para._element.addprevious(new_elem)
            new_para = Paragraph(new_elem, para._parent)
            run = new_para.add_run()
            run.add_picture(img_path, width=Inches(5.0))
            new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_elem = etree.SubElement(para._element.getparent(), qn('w:p'))
            para._element.addprevious(cap_elem)
            cap_para = Paragraph(cap_elem, para._parent)
            cap_run = cap_para.add_run('Fig 1.1: Traditional Blacklist vs ML-Based URL Detection')
            cap_run.font.size = Pt(10); cap_run.font.italic = True; cap_run.font.name = 'Times New Roman'
            cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            blank_elem = etree.SubElement(para._element.getparent(), qn('w:p'))
            para._element.addprevious(blank_elem)
            fig_1_1_done = True
            inserted += 1
            print(f"\n  Inserted Fig 1.1 before paragraph [{i}]")
        break

doc.save(DOC_PATH)
print(f"\nFig 1.1: {fig_1_1_done}")
print(f"\nTotal figures inserted: {inserted}")
print(f"Document saved: {DOC_PATH}")
