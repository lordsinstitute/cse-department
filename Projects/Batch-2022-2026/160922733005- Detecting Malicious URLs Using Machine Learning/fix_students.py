"""Replace student details in C12 report."""

from docx import Document

DOC_PATH = '/Users/shoukathali/lord-major-projects/IV-C Projects/C12/Detecting_Malicious_URLs_Using_ML_Major_Project_Report.docx'

OLD = [
    ('Muhammad Aasim Uz Zaman', '160922733020'),
    ('Syed Altamash Uddin Siddiqui', '160922733032'),
    ('Nawaz Khan', '160922733037'),
    ('Faiz Ur Rahman', '160922733049'),
]

NEW = [
    ('Mohammed Masood Ullah', '160922733005'),
    ('Fouzan Mohammed Khan', '160922733017'),
    ('Muaaz Hamad', '160922733009'),
    ('Shoaib Mohammed Basheer', '160922733002'),
]

doc = Document(DOC_PATH)

# Fix tables 0, 5, 7
for ti in [0, 5, 7]:
    tbl = doc.tables[ti]
    for ri in range(4):
        for run in tbl.rows[ri].cells[0].paragraphs[0].runs:
            if OLD[ri][0] in run.text:
                run.text = run.text.replace(OLD[ri][0], NEW[ri][0])
        for run in tbl.rows[ri].cells[1].paragraphs[0].runs:
            if OLD[ri][1] in run.text:
                run.text = run.text.replace(OLD[ri][1], NEW[ri][1])
    print(f"  Table[{ti}]: Updated 4 student names/rolls")

# Fix certificate paragraph
old_str = 'Muhammad Aasim Uz Zaman (160922733020), Syed Altamash Uddin Siddiqui (160922733032), Nawaz Khan (160922733037), and Faiz Ur Rahman (160922733049)'
new_str = 'Mohammed Masood Ullah (160922733005), Fouzan Mohammed Khan (160922733017), Muaaz Hamad (160922733009), and Shoaib Mohammed Basheer (160922733002)'

for i, p in enumerate(doc.paragraphs):
    if '160922733020' in p.text and '160922733032' in p.text:
        full_text = p.text
        new_text = full_text.replace(old_str, new_str)
        if new_text != full_text:
            for run in p.runs:
                run.text = ''
            p.runs[0].text = new_text
            print(f"  Paragraph[{i}]: Updated certificate text")
        break

# Catch any remaining old names/rolls in other paragraphs
for i, p in enumerate(doc.paragraphs):
    for run in p.runs:
        for oi in range(4):
            if OLD[oi][0] in run.text:
                run.text = run.text.replace(OLD[oi][0], NEW[oi][0])
            if OLD[oi][1] in run.text:
                run.text = run.text.replace(OLD[oi][1], NEW[oi][1])

doc.save(DOC_PATH)
print(f"\nDone. Saved: {DOC_PATH}")
